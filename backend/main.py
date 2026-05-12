"""File Knowledge Base - Backend API with Auth"""
import json
import uuid
import hashlib
import re
import os
import secrets
import sqlite3
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional
from collections import Counter

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from pydantic import BaseModel
import pandas as pd
from openpyxl import load_workbook, Workbook
from docx import Document
from io import BytesIO
import shutil

app = FastAPI(title="File Knowledge Base")


def _parse_cors_origins() -> list[str]:
    raw = os.environ.get("CORS_ORIGINS", "")
    if not raw.strip():
        return [
            "http://localhost:8000",
            "http://127.0.0.1:8000",
            "http://localhost:5173",
            "http://127.0.0.1:5173",
        ]
    return [o.strip() for o in raw.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_parse_cors_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).parent.parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
UPLOAD_VERSIONS_DIR = UPLOAD_DIR / "versions"
UPLOAD_VERSIONS_DIR.mkdir(exist_ok=True)

DB_FILE = BASE_DIR / "kb.db"

# Quota Constants
USER_QUOTA_DEFAULT = 2 * 1024 * 1024 * 1024  # 2GB
WORKSPACE_QUOTA_DEFAULT = 100 * 1024 * 1024 * 1024  # 100GB

# Legacy JSON file paths (for one-time migration)
JSON_DB = BASE_DIR / "kb_index.json"
JSON_USERS = BASE_DIR / "backend" / "kb_users.json"
JSON_SESSIONS = BASE_DIR / "backend" / "kb_sessions.json"

# ============ SQLite Database ============

def get_db():
    conn = sqlite3.connect(str(DB_FILE))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password_hash TEXT NOT NULL,
            name TEXT NOT NULL DEFAULT '',
            role TEXT NOT NULL DEFAULT 'member',
            status TEXT NOT NULL DEFAULT 'active',
            created_at TEXT NOT NULL,
            bio TEXT NOT NULL DEFAULT '',
            avatar_data TEXT NOT NULL DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS sessions (
            token TEXT PRIMARY KEY,
            username TEXT NOT NULL,
            expires TEXT NOT NULL,
            FOREIGN KEY (username) REFERENCES users(username)
        );
        CREATE TABLE IF NOT EXISTS files (
            id TEXT PRIMARY KEY,
            filename TEXT NOT NULL,
            size INTEGER NOT NULL DEFAULT 0,
            type TEXT NOT NULL,
            folder TEXT NOT NULL DEFAULT '',
            view_count INTEGER NOT NULL DEFAULT 0,
            uploaded_at TEXT NOT NULL,
            uploaded_by TEXT NOT NULL,
            analysis TEXT NOT NULL DEFAULT '{}',
            FOREIGN KEY (uploaded_by) REFERENCES users(username)
        );
        CREATE TABLE IF NOT EXISTS folders (
            folder TEXT PRIMARY KEY
        );
        CREATE TABLE IF NOT EXISTS file_versions (
            id TEXT PRIMARY KEY,
            file_id TEXT NOT NULL,
            version_num INTEGER NOT NULL DEFAULT 1,
            size INTEGER NOT NULL,
            saved_at TEXT NOT NULL,
            saved_by TEXT NOT NULL,
            FOREIGN KEY (file_id) REFERENCES files(id)
        );
        CREATE TABLE IF NOT EXISTS file_locks (
            file_id TEXT PRIMARY KEY,
            username TEXT NOT NULL,
            name TEXT NOT NULL,
            last_heartbeat TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS file_favorites (
            username TEXT NOT NULL,
            file_id TEXT NOT NULL,
            created_at TEXT NOT NULL,
            PRIMARY KEY (username, file_id),
            FOREIGN KEY (username) REFERENCES users(username),
            FOREIGN KEY (file_id) REFERENCES files(id)
        );
        CREATE TABLE IF NOT EXISTS announcements (
            id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            content TEXT NOT NULL DEFAULT '',
            created_by TEXT NOT NULL,
            created_at TEXT NOT NULL,
            is_active INTEGER NOT NULL DEFAULT 1
        );
        CREATE TABLE IF NOT EXISTS admin_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            actor TEXT NOT NULL,
            action TEXT NOT NULL,
            target TEXT NOT NULL,
            details TEXT DEFAULT '',
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS announcement_reads (
            announcement_id TEXT NOT NULL,
            username TEXT NOT NULL,
            read_at TEXT NOT NULL,
            PRIMARY KEY (announcement_id, username),
            FOREIGN KEY (announcement_id) REFERENCES announcements(id),
            FOREIGN KEY (username) REFERENCES users(username)
        );
        CREATE TABLE IF NOT EXISTS workspaces (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            description TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            created_by TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS workspace_members (
            workspace_id TEXT NOT NULL,
            username TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'member',
            joined_at TEXT NOT NULL,
            PRIMARY KEY (workspace_id, username)
        );
        CREATE TABLE IF NOT EXISTS file_permissions (
            file_id TEXT PRIMARY KEY,
            visibility TEXT NOT NULL DEFAULT 'workspace',
            allow_view INTEGER NOT NULL DEFAULT 1,
            allow_edit INTEGER NOT NULL DEFAULT 0,
            allow_delete INTEGER NOT NULL DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS file_transfers (
            id TEXT PRIMARY KEY,
            file_id TEXT NOT NULL,
            sender TEXT NOT NULL,
            recipient TEXT NOT NULL,
            message TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'pending',
            created_at TEXT NOT NULL,
            resolved_at TEXT NOT NULL DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS file_links (
            id TEXT PRIMARY KEY,
            file_id TEXT NOT NULL,
            related_file_id TEXT NOT NULL,
            relation_label TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            created_by TEXT NOT NULL,
            UNIQUE(file_id, related_file_id)
        );
        CREATE TABLE IF NOT EXISTS file_link_logs (
            id TEXT PRIMARY KEY,
            link_id TEXT NOT NULL,
            action TEXT NOT NULL,
            file_id TEXT NOT NULL,
            related_file_id TEXT NOT NULL,
            relation_label TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            created_by TEXT NOT NULL
        );
    """)
    conn.commit()
    conn.close()

def migrate_versions():
    """Ensure file_versions and file_locks tables exist, and migrate files table."""
    conn = get_db()
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS file_versions (
            id TEXT PRIMARY KEY,
            file_id TEXT NOT NULL,
            version_num INTEGER NOT NULL DEFAULT 1,
            size INTEGER NOT NULL,
            saved_at TEXT NOT NULL,
            saved_by TEXT NOT NULL,
            FOREIGN KEY (file_id) REFERENCES files(id)
        );
        CREATE TABLE IF NOT EXISTS file_locks (
            file_id TEXT PRIMARY KEY,
            username TEXT NOT NULL,
            name TEXT NOT NULL,
            last_heartbeat TEXT NOT NULL
        );
    """)
    # Add last_modified_at / last_modified_by if not already present
    for col, definition in [("last_modified_at", "TEXT"), ("last_modified_by", "TEXT")]:
        try:
            c.execute(f"ALTER TABLE files ADD COLUMN {col} {definition}")
        except Exception:
            pass
    # FTS5 full-text search tables
    try:
        c.execute("""
            CREATE VIRTUAL TABLE IF NOT EXISTS files_fts USING fts5(
                filename,
                content,
                tokenize='unicode61 remove_diacritics 1'
            )
        """)
        c.execute("""
            CREATE TABLE IF NOT EXISTS fts_file_map (
                file_id TEXT PRIMARY KEY,
                fts_rowid INTEGER NOT NULL
            )
        """)
    except Exception as e:
        print(f"FTS5 init warning: {e}")
    conn.commit()
    conn.close()

# ============ FTS5 Full-Text Search Helpers ============

def fts_extract_content(file_id: str, file_ext: str) -> str:
    """Extract plain text from a file for FTS indexing (max 100 KB)."""
    file_path = UPLOAD_DIR / f"{file_id}{file_ext}"
    if not file_path.exists():
        return ""
    ext = file_ext.lower()
    try:
        if ext in ('.txt', '.md', '.json', '.xml', '.yaml', '.yml',
                   '.py', '.sql', '.sh', '.bat', '.css', '.js',
                   '.csv', '.tsv', '.log'):
            return file_path.read_text(encoding='utf-8', errors='ignore')[:100_000]
        if ext == '.html':
            raw = file_path.read_text(encoding='utf-8', errors='ignore')
            text = re.sub(r'<style[^>]*>.*?</style>', ' ', raw, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<script[^>]*>.*?</script>', ' ', text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<[^>]+>', ' ', text)
            return re.sub(r'\s+', ' ', text).strip()[:100_000]
        if ext == '.docx':
            doc = Document(str(file_path))
            return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())[:100_000]
        if ext in ('.xlsx', '.xls'):
            wb = load_workbook(str(file_path), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    parts.append(' '.join(str(cell) for cell in row if cell is not None))
            wb.close()
            return '\n'.join(parts)[:100_000]
    except Exception as e:
        print(f"FTS extract error {file_id}{file_ext}: {e}")
    return ""


def fts_index_file(conn, file_id: str, filename: str, file_ext: str):
    """Insert or replace a file's FTS index entry."""
    content = fts_extract_content(file_id, file_ext)
    try:
        row = conn.execute(
            "SELECT fts_rowid FROM fts_file_map WHERE file_id = ?", (file_id,)
        ).fetchone()
        if row:
            conn.execute("DELETE FROM files_fts WHERE rowid = ?", (row["fts_rowid"],))
            conn.execute("DELETE FROM fts_file_map WHERE file_id = ?", (file_id,))
        cur = conn.execute(
            "INSERT INTO files_fts (filename, content) VALUES (?, ?)", (filename, content)
        )
        conn.execute(
            "INSERT INTO fts_file_map (file_id, fts_rowid) VALUES (?, ?)",
            (file_id, cur.lastrowid),
        )
    except Exception as e:
        print(f"FTS index error {file_id}: {e}")


def fts_delete_file(conn, file_id: str):
    """Remove a file's FTS index entry."""
    try:
        row = conn.execute(
            "SELECT fts_rowid FROM fts_file_map WHERE file_id = ?", (file_id,)
        ).fetchone()
        if row:
            conn.execute("DELETE FROM files_fts WHERE rowid = ?", (row["fts_rowid"],))
            conn.execute("DELETE FROM fts_file_map WHERE file_id = ?", (file_id,))
    except Exception as e:
        print(f"FTS delete error {file_id}: {e}")


def build_fts_query(raw: str) -> str | None:
    """Convert a user search string to a safe FTS5 MATCH expression."""
    q = re.sub(r'["\*\(\)\:\.\^]', ' ', raw.strip())
    tokens = [t for t in q.split() if t]
    if not tokens:
        return None
    parts = []
    for token in tokens:
        if re.search(r'[涓€-榭縘', token):
            # For CJK: split into chars and use FTS5 phrase syntax for adjacency
            chars = ' '.join(list(token))
            parts.append(f'"{chars}"')
        else:
            parts.append(f'"{token}"*')
    return ' '.join(parts)


def build_fts_index():
    """At startup, index all existing files not yet in FTS."""
    conn = get_db()
    try:
        all_files = conn.execute("SELECT id, filename, type FROM files").fetchall()
        indexed = {r["file_id"] for r in conn.execute("SELECT file_id FROM fts_file_map").fetchall()}
        count = 0
        for f in all_files:
            if f["id"] not in indexed:
                fts_index_file(conn, f["id"], f["filename"], f["type"])
                count += 1
        conn.commit()
        if count:
            print(f"FTS: indexed {count} existing files")
    except Exception as e:
        print(f"FTS build_index error: {e}")
    finally:
        conn.close()


# ============ Lock & Draft Helpers ============
UPLOAD_DRAFTS_DIR = UPLOAD_DIR / "drafts"
UPLOAD_DRAFTS_DIR.mkdir(exist_ok=True)

def get_file_lock_info(file_id: str):
    """Check if file is locked and by whom. Returns (is_locked, locker_info)."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM file_locks WHERE file_id = ?", (file_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        return False, None
    
    # Check expiry (30 seconds)
    try:
        last_hb = datetime.fromisoformat(row["last_heartbeat"])
        if datetime.now() - last_hb > timedelta(seconds=30):
            # Lock expired
            return False, None
    except Exception:
        return False, None
    
    return True, {"username": row["username"], "name": row["name"]}

def migrate_from_json():
    """One-time migration: import data from legacy JSON files if DB doesn't exist yet."""
    needs_migration = not DB_FILE.exists()

    init_db()

    if not needs_migration:
        return

    conn = get_db()
    c = conn.cursor()

    # Migrate users
    if JSON_USERS.exists():
        with open(JSON_USERS, encoding="utf-8") as f:
            users_data = json.load(f)
        for u in users_data.get("users", []):
            if "status" not in u:
                u["status"] = "active"
            c.execute("INSERT OR IGNORE INTO users (username, password_hash, name, role, status, created_at) VALUES (?, ?, ?, ?, ?, ?)",
                      (u["username"], u["password_hash"], u.get("name", u["username"]), u.get("role", "member"), u.get("status", "active"), u.get("created_at", datetime.now().isoformat())))

    # Migrate sessions
    if JSON_SESSIONS.exists():
        with open(JSON_SESSIONS, encoding="utf-8") as f:
            sessions_data = json.load(f)
        for token, sess in sessions_data.get("sessions", {}).items():
            c.execute("INSERT OR IGNORE INTO sessions (token, username, expires) VALUES (?, ?, ?)",
                      (token, sess["username"], sess["expires"]))

    # Migrate files and folders
    if JSON_DB.exists():
        with open(JSON_DB, encoding="utf-8") as f:
            file_data = json.load(f)
        for fid, rec in file_data.get("files", {}).items():
            if "folder" not in rec:
                rec["folder"] = ""
            if "view_count" not in rec:
                rec["view_count"] = 0
            analysis_json = json.dumps(rec.get("analysis", {}), ensure_ascii=False)
            c.execute("INSERT OR IGNORE INTO files (id, filename, size, type, folder, view_count, uploaded_at, uploaded_by, analysis) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                      (fid, rec["filename"], rec.get("size", 0), rec.get("type", ""), rec["folder"], rec["view_count"], rec.get("uploaded_at", ""), rec.get("uploaded_by", ""), analysis_json))
        for folder in file_data.get("folders", []):
            c.execute("INSERT OR IGNORE INTO folders (folder) VALUES (?)", (folder,))

    conn.commit()
    conn.close()

migrate_from_json()

# ============ Helpers ============

def _load_json(path, default=None):
    if default is None:
        default = {}
    if path.exists():
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return default

def hash_pw(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def is_strong_password(pw: str) -> bool:
    """Check if password contains letters, numbers and is at least 8 chars."""
    if len(pw) < 8:
        return False
    if not re.search(r"[a-zA-Z]", pw):
        return False
    if not re.search(r"[0-9]", pw):
        return False
    return True

# ============ Auth ============

def get_user(authorization: Optional[str] = Header(None)):
    """楠岃瘉 token锛岃繑鍥炵敤鎴蜂俊鎭垨鎶?401"""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(401, "鏈櫥褰?")
    token = authorization[7:]
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT s.username, s.expires, u.role, u.status, u.name, u.bio, u.avatar_data FROM sessions s JOIN users u ON s.username = u.username WHERE s.token = ?", (token,))
    row = c.fetchone()
    conn.close()
    if not row:
        raise HTTPException(401, "鐧诲綍宸茶繃鏈?")
    if datetime.fromisoformat(row["expires"]) < datetime.now():
        conn = get_db()
        conn.execute("DELETE FROM sessions WHERE token = ?", (token,))
        conn.commit()
        conn.close()
        raise HTTPException(401, "鐧诲綍宸茶繃鏈?")
    if row["status"] != "active":
        raise HTTPException(403, "璐﹀彿寰呭鎵癸紝璇疯仈绯荤鐞嗗憳")
    return {
        "username": row["username"], "name": row["name"], "role": row["role"],
        "status": row["status"], "bio": row["bio"] or "", "avatar_data": row["avatar_data"] or "",
    }

def get_admin_user(authorization: Optional[str] = Header(None)):
    """楠岃瘉 token 涓斾负 admin 瑙掕壊"""
    user = get_user(authorization)
    if user.get("role") != "admin":
        raise HTTPException(403, "闇€瑕佺鐞嗗憳鏉冮檺")
    return user

def require_auth(func):
    """瑁呴グ鍣細缁欒矾鐢卞嚱鏁板姞閴存潈"""
    from functools import wraps
    @wraps(func)
    async def wrapper(authorization: Optional[str] = Header(None), **kwargs):
        user = get_user(authorization)
        kwargs["current_user"] = user
        return await func(**kwargs)
    return wrapper

def require_admin(func):
    from functools import wraps
    @wraps(func)
    async def wrapper(authorization: Optional[str] = Header(None), **kwargs):
        user = get_admin_user(authorization)
        kwargs["current_user"] = user
        return await func(**kwargs)
    return wrapper

# ============ Default Admin ============

def init_default_admin():
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) as cnt FROM users")
    count = c.fetchone()["cnt"]
    if count == 0:
        init_pw = os.environ.get("ADMIN_INIT_PASSWORD", "").strip()
        if not init_pw:
            init_pw = secrets.token_urlsafe(12)
            print(f"[SECURITY] Generated initial admin password: {init_pw}")
        c.execute("INSERT INTO users (username, password_hash, name, role, status, created_at) VALUES (?, ?, ?, ?, ?, ?)",
                  ("admin", hash_pw(init_pw), "管理员", "admin", "active", datetime.now().isoformat()))
    conn.commit()
    conn.close()

# ============ Permission & Profile Migration ============

def migrate_permissions():
    """Add bio, avatar, permissions columns to users table."""
    conn = get_db()
    c = conn.cursor()

    for col, default in [
        ("bio", "''"),
        ("avatar_data", "''"),
        ("permissions", "''"),
    ]:
        try:
            c.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT NOT NULL DEFAULT {default}")
        except Exception:
            pass  # column already exists

    conn.commit()
    conn.close()

def migrate_workspaces():
    """Add workspace tables and workspace_id column to files."""
    conn = get_db()
    c = conn.cursor()
    c.executescript(f"""
        CREATE TABLE IF NOT EXISTS workspaces (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            description TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            created_by TEXT NOT NULL,
            quota_limit INTEGER NOT NULL DEFAULT {WORKSPACE_QUOTA_DEFAULT},
            quota_used INTEGER NOT NULL DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS workspace_members (
            workspace_id TEXT NOT NULL,
            username TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'member',
            joined_at TEXT NOT NULL,
            quota_limit INTEGER NOT NULL DEFAULT {USER_QUOTA_DEFAULT},
            quota_used INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY (workspace_id, username)
        );
        CREATE TABLE IF NOT EXISTS file_permissions (
            file_id TEXT PRIMARY KEY,
            visibility TEXT NOT NULL DEFAULT 'workspace',
            allow_view INTEGER NOT NULL DEFAULT 1,
            allow_edit INTEGER NOT NULL DEFAULT 0,
            allow_delete INTEGER NOT NULL DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS workspace_join_requests (
            id TEXT PRIMARY KEY,
            workspace_id TEXT NOT NULL,
            username TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            requested_at TEXT NOT NULL,
            reviewed_by TEXT NOT NULL DEFAULT '',
            reviewed_at TEXT NOT NULL DEFAULT '',
            UNIQUE (workspace_id, username)
        );
    """)
    # Add columns if they don't exist (for existing tables)
    for table, col, dtype, default in [
        ("workspaces", "quota_limit", "INTEGER", WORKSPACE_QUOTA_DEFAULT),
        ("workspaces", "quota_used", "INTEGER", 0),
        ("workspace_members", "quota_limit", "INTEGER", USER_QUOTA_DEFAULT),
        ("workspace_members", "quota_used", "INTEGER", 0),
    ]:
        try:
            c.execute(f"ALTER TABLE {table} ADD COLUMN {col} {dtype} NOT NULL DEFAULT {default}")
        except Exception:
            pass

    try:
        c.execute("ALTER TABLE files ADD COLUMN workspace_id TEXT NOT NULL DEFAULT ''")
    except Exception:
        pass
    conn.commit()
    conn.close()

def check_quota(username: str, workspace_id: str, additional_size: int, conn):
    """Check if adding additional_size will exceed user or workspace quota."""
    # Check if user is super admin
    user = conn.execute("SELECT role FROM users WHERE username = ?", (username,)).fetchone()
    if user and user["role"] == "admin":
        return

    if not workspace_id:
        raise HTTPException(403, "涓婁紶鏂囦欢蹇呴』鎸囧畾鎵€灞炵殑宸ヤ綔绌洪棿锛屾湭鍔犲叆宸ヤ綔绌洪棿鍓嶆病鏈夊瓨鍌ㄩ搴?")

    # Check workspace quota
    ws = conn.execute("SELECT quota_limit, quota_used FROM workspaces WHERE id = ?", (workspace_id,)).fetchone()
    if not ws:
        raise HTTPException(404, "宸ヤ綔绌洪棿涓嶅瓨鍦?")
    
    if ws["quota_used"] + additional_size > ws["quota_limit"]:
        limit_gb = ws["quota_limit"] / (1024**3)
        raise HTTPException(400, f"宸ヤ綔绌洪棿瀛樺偍棰濆害涓嶈冻 (鏈€澶? {limit_gb:.1f}GB)")

    # Check user quota in workspace
    member = conn.execute("SELECT quota_limit, quota_used FROM workspace_members WHERE workspace_id = ? AND username = ?", (workspace_id, username)).fetchone()
    if not member:
        raise HTTPException(403, "鎮ㄤ笉鏄宸ヤ綔绌洪棿鐨勬垚鍛?")
    
    if member["quota_used"] + additional_size > member["quota_limit"]:
        limit_mb = member["quota_limit"] / (1024**2)
        raise HTTPException(400, f"鎮ㄧ殑涓汉瀛樺偍棰濆害涓嶈冻 (鏈€澶? {limit_mb:.1f}MB)")

def update_quota(username: str, workspace_id: str, size_change: int, conn):
    """Update quota_used for user and workspace."""
    if not workspace_id:
        return
    conn.execute("UPDATE workspace_members SET quota_used = quota_used + ? WHERE workspace_id = ? AND username = ?", (size_change, workspace_id, username))
    conn.execute("UPDATE workspaces SET quota_used = quota_used + ? WHERE id = ?", (size_change, workspace_id))

def recalculate_quotas():
    """Recalculate all quota_used based on current files and versions."""
    conn = get_db()
    # Reset
    conn.execute("UPDATE workspace_members SET quota_used = 0")
    conn.execute("UPDATE workspaces SET quota_used = 0")
    
    # Files
    files = conn.execute("SELECT workspace_id, uploaded_by, size FROM files WHERE workspace_id != ''").fetchall()
    for f in files:
        update_quota(f["uploaded_by"], f["workspace_id"], f["size"], conn)
        
    # Versions
    versions = conn.execute("""
        SELECT f.workspace_id, f.uploaded_by, v.size 
        FROM file_versions v 
        JOIN files f ON v.file_id = f.id 
        WHERE f.workspace_id != ''
    """).fetchall()
    for v in versions:
        update_quota(v["uploaded_by"], v["workspace_id"], v["size"], conn)
    
    conn.commit()
    conn.close()

def migrate_file_transfers():
    """Ensure file_transfers table exists."""
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS file_transfers (
            id TEXT PRIMARY KEY,
            file_id TEXT NOT NULL,
            sender TEXT NOT NULL,
            recipient TEXT NOT NULL,
            message TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'pending',
            created_at TEXT NOT NULL,
            resolved_at TEXT NOT NULL DEFAULT ''
        )
    """)
    conn.commit()
    conn.close()

migrate_permissions()
migrate_versions()
migrate_workspaces()
migrate_file_transfers()
recalculate_quotas()
build_fts_index()
init_default_admin()

DEFAULT_MEMBER_PERMS = {
    "can_upload": True,
    "can_delete_own": True,
    "can_delete_any": False,
    "can_edit_own": True,
    "can_edit_others": False,
    "can_create_folder": False,
    "can_view_report": False,
}

ADMIN_PERMS = {
    "can_upload": True,
    "can_delete_own": True,
    "can_delete_any": True,
    "can_edit_own": True,
    "can_edit_others": True,
    "can_create_folder": True,
    "can_view_report": True,
}

def get_permissions(username: str) -> dict:
    conn = get_db()
    row = conn.execute("SELECT role, permissions FROM users WHERE username = ?", (username,)).fetchone()
    conn.close()
    if not row:
        return {}
    if row["role"] == "admin":
        return dict(ADMIN_PERMS)
    perms_str = row["permissions"] or ""
    if perms_str:
        try:
            loaded = json.loads(perms_str)
            merged = dict(DEFAULT_MEMBER_PERMS)
            if isinstance(loaded, dict):
                merged.update(loaded)
            return merged
        except Exception:
            pass
    return dict(DEFAULT_MEMBER_PERMS)

def check_perm(user: dict, perm: str, *, is_owner: bool = False):
    """Raise 403 if user lacks permission. Admin always passes."""
    if user["role"] == "admin":
        return
    perms = get_permissions(user["username"])
    if not perms.get(perm, False):
        raise HTTPException(403, "鏃犳潈闄愭墽琛屾鎿嶄綔")
    if not is_owner and perm in ("can_delete_any", "can_edit_others"):
        raise HTTPException(403, "鏃犳潈闄愭搷浣滀粬浜烘枃浠?")

def log_admin_action(actor: str, action: str, target: str, details: str = ""):
    conn = get_db()
    conn.execute(
        "INSERT INTO admin_logs (actor, action, target, details, created_at) VALUES (?, ?, ?, ?, ?)",
        (actor, action, target, details, datetime.now().isoformat())
    )
    conn.commit()
    conn.close()

# ============ Auth API ============

@app.get("/api/admin/logs")
async def get_admin_logs(current_user: dict = Depends(get_admin_user)):
    conn = get_db()
    rows = conn.execute("SELECT * FROM admin_logs ORDER BY created_at DESC LIMIT 100").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/api/users/batch-approve")
async def batch_approve_users(req: dict, current_user: dict = Depends(get_admin_user)):
    usernames = req.get("usernames", [])
    if not usernames:
        return {"ok": True}
    conn = get_db()
    placeholders = ",".join(["?"] * len(usernames))
    conn.execute(f"UPDATE users SET status = 'active' WHERE username IN ({placeholders}) AND status = 'pending'", usernames)
    conn.commit()
    conn.close()
    log_admin_action(current_user["username"], "batch_approve", f"{len(usernames)} users", ",".join(usernames))
    return {"ok": True}

@app.post("/api/users/batch-delete")
async def batch_delete_users(req: dict, current_user: dict = Depends(get_admin_user)):
    usernames = req.get("usernames", [])
    if not usernames:
        return {"ok": True}
    # Protect self
    usernames = [u for u in usernames if u != current_user["username"]]
    conn = get_db()
    placeholders = ",".join(["?"] * len(usernames))
    conn.execute(f"DELETE FROM users WHERE username IN ({placeholders})", usernames)
    conn.commit()
    conn.close()
    log_admin_action(current_user["username"], "batch_delete", f"{len(usernames)} users", ",".join(usernames))
    return {"ok": True}

class LoginRequest(BaseModel):
    username: str
    password: str

class RegisterRequest(BaseModel):
    username: str
    password: str
    name: str = ""

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str

class AnnouncementCreate(BaseModel):
    title: str
    content: str = ""

class AnnouncementUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    is_active: Optional[bool] = None

class FileCreateRequest(BaseModel):
    filename: str
    file_type: str
    content: str = ""
    folder: str = ""
    workspace_id: str = ""
    visibility: str = "workspace"
    allow_view: bool = True
    allow_edit: bool = False
    allow_delete: bool = False

class FileTransferRequest(BaseModel):
    recipient: str
    message: str = ""

class FileTransferAcceptRequest(BaseModel):
    workspace_ids: list = []
    visibility: str = "workspace"
    allow_view: bool = True
    allow_edit: bool = False
    allow_delete: bool = False

class FileUpdateRequest(BaseModel):
    content: str

class WorkspaceCreate(BaseModel):
    name: str
    description: str = ""
    quota_limit: int = WORKSPACE_QUOTA_DEFAULT

class WorkspaceUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    quota_limit: Optional[int] = None

class WorkspaceMemberUpdate(BaseModel):
    username: str
    role: str = "member"
    quota_limit: int = USER_QUOTA_DEFAULT

class FilePermissionsUpdate(BaseModel):
    visibility: str = "workspace"
    allow_view: bool = True
    allow_edit: bool = False
    allow_delete: bool = False

class FileLinkCreateRequest(BaseModel):
    file_id: str
    related_file_id: str
    relation_label: str = ""

class QaAskRequest(BaseModel):
    question: str
    messages: Optional[list] = []  # 鍘嗗彶瀵硅瘽 [{"role": "user"|"assistant", "content": str}]

class SettingsUpdate(BaseModel):
    llm_api_key: Optional[str] = ""
    llm_model: Optional[str] = ""
    llm_base_url: Optional[str] = ""
    llm_configs: Optional[str] = "[]"
    active_llm_index: Optional[int] = 0

class LlmTestRequest(BaseModel):
    llm_api_key: str
    llm_model: str
    llm_base_url: str

@app.post("/api/register")
async def register(req: RegisterRequest, authorization: Optional[str] = Header(None)):
    if not is_strong_password(req.password):
        raise HTTPException(400, "瀵嗙爜寮哄害涓嶈冻锛氬繀椤诲寘鍚瓧姣嶅拰鏁板瓧锛屼笖涓嶅皯浜?8 浣?")
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) as cnt FROM users")
    is_first = c.fetchone()["cnt"] == 0

    c.execute("SELECT COUNT(*) as cnt FROM users WHERE username = ?", (req.username,))
    if c.fetchone()["cnt"] > 0:
        conn.close()
        raise HTTPException(400, "鐢ㄦ埛鍚嶅凡瀛樺湪")

    # 妫€鏌ユ槸鍚︽槸绠＄悊鍛樹唬鍒?
    is_admin_create = False
    if not is_first and authorization and authorization.startswith("Bearer "):
        tkn = authorization[7:]
        c.execute("SELECT s.username, u.role FROM sessions s JOIN users u ON s.username = u.username WHERE s.token = ? AND s.expires > ?", (tkn, datetime.now().isoformat()))
        sess = c.fetchone()
        if sess and sess["role"] == "admin":
            is_admin_create = True

    if is_first:
        role = "admin"
        status = "active"
    elif is_admin_create:
        role = "member"
        status = "active"
    else:
        role = "member"
        status = "pending"

    now = datetime.now().isoformat()
    c.execute("INSERT INTO users (username, password_hash, name, role, status, created_at) VALUES (?, ?, ?, ?, ?, ?)",
              (req.username, hash_pw(req.password), req.name or req.username, role, status, now))
    conn.commit()

    result = {"username": req.username, "name": req.name or req.username, "role": role, "status": status}
    if is_first:
        token = uuid.uuid4().hex
        c.execute("INSERT INTO sessions (token, username, expires) VALUES (?, ?, ?)",
                  (token, req.username, (datetime.now() + timedelta(hours=24)).isoformat()))
        conn.commit()
        conn.close()
        return {"token": token, **result}
    elif is_admin_create:
        conn.close()
        return {"ok": True, **result}
    else:
        conn.close()
        return {"status": "pending", **result, "message": "注册成功，请等待管理员审批"}

@app.post("/api/login")
async def login(req: LoginRequest):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT username, name, role, status FROM users WHERE username = ? AND password_hash = ?", (req.username, hash_pw(req.password)))
    u = c.fetchone()
    if not u:
        conn.close()
        raise HTTPException(401, "鐢ㄦ埛鍚嶆垨瀵嗙爜閿欒")
    if u["status"] != "active":
        conn.close()
        raise HTTPException(403, "璐﹀彿寰呭鎵癸紝璇疯仈绯荤鐞嗗憳")
    token = uuid.uuid4().hex
    c.execute("INSERT INTO sessions (token, username, expires) VALUES (?, ?, ?)",
              (token, u["username"], (datetime.now() + timedelta(hours=24)).isoformat()))
    conn.commit()
    conn.close()
    return {"token": token, "username": u["username"], "name": u["name"], "role": u["role"], "status": u["status"]}

@app.post("/api/logout")
async def logout(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        return {"ok": True}
    token = authorization[7:]
    conn = get_db()
    conn.execute("DELETE FROM sessions WHERE token = ?", (token,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/api/me")
async def me(authorization: Optional[str] = Header(None)):
    user = get_user(authorization)
    return user

class ProfileUpdateRequest(BaseModel):
    name: str = ""
    bio: str = ""
    avatar_data: str = ""

@app.post("/api/profile")
async def update_profile(req: ProfileUpdateRequest, authorization: Optional[str] = Header(None)):
    user = get_user(authorization)
    conn = get_db()
    conn.execute("UPDATE users SET name = ?, bio = ?, avatar_data = ? WHERE username = ?",
                 (req.name or user["username"], req.bio, req.avatar_data, user["username"]))
    conn.commit()
    conn.close()
    return {"ok": True}
    conn.close()
    return {"ok": True}

@app.post("/api/change-password")
async def change_password(req: ChangePasswordRequest, authorization: Optional[str] = Header(None)):
    if not is_strong_password(req.new_password):
        raise HTTPException(400, "鏂板瘑鐮佸己搴︿笉瓒筹細蹇呴』鍖呭惈瀛楁瘝鍜屾暟瀛楋紝涓斾笉灏戜簬 8 浣?")
    user = get_user(authorization)
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT password_hash FROM users WHERE username = ?", (user["username"],))
    row = c.fetchone()
    if not row or row[0] != hash_pw(req.old_password):
        conn.close()
        raise HTTPException(400, "鏃у瘑鐮佷笉姝ｇ‘")
    conn.execute("UPDATE users SET password_hash = ? WHERE username = ?",
                 (hash_pw(req.new_password), user["username"]))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/api/me/permissions")
async def my_permissions(authorization: Optional[str] = Header(None)):
    user = get_user(authorization)
    return get_permissions(user["username"])

# ============ Announcement API ============

@app.post("/api/announcements")
async def create_announcement(req: AnnouncementCreate, current_user: dict = Depends(get_admin_user)):
    aid = uuid.uuid4().hex[:12]
    now = datetime.now().isoformat()
    conn = get_db()
    conn.execute(
        "INSERT INTO announcements (id, title, content, created_by, created_at, is_active) VALUES (?, ?, ?, ?, ?, 1)",
        (aid, req.title, req.content, current_user["username"], now)
    )
    conn.commit()
    conn.close()
    return {"ok": True, "id": aid}

@app.get("/api/announcements")
async def list_announcements(current_user: dict = Depends(get_user)):
    conn = get_db()
    if current_user["role"] == "admin":
        rows = conn.execute(
            "SELECT a.*, u.name as creator_name FROM announcements a LEFT JOIN users u ON a.created_by = u.username ORDER BY a.created_at DESC"
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT a.*, u.name as creator_name FROM announcements a LEFT JOIN users u ON a.created_by = u.username WHERE a.is_active = 1 ORDER BY a.created_at DESC"
        ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/api/announcements/unread")
async def get_unread_announcements(current_user: dict = Depends(get_user)):
    conn = get_db()
    rows = conn.execute("""
        SELECT a.*, u.name as creator_name
        FROM announcements a
        LEFT JOIN users u ON a.created_by = u.username
        WHERE a.is_active = 1
          AND a.id NOT IN (
              SELECT announcement_id FROM announcement_reads WHERE username = ?
          )
        ORDER BY a.created_at DESC
    """, (current_user["username"],)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.patch("/api/announcements/{announcement_id}")
async def update_announcement(announcement_id: str, req: AnnouncementUpdate, current_user: dict = Depends(get_admin_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id FROM announcements WHERE id = ?", (announcement_id,))
    if not c.fetchone():
        conn.close()
        raise HTTPException(404, "鍏憡涓嶅瓨鍦?")
    updates = []
    params = []
    if req.title is not None:
        updates.append("title = ?")
        params.append(req.title)
    if req.content is not None:
        updates.append("content = ?")
        params.append(req.content)
    if req.is_active is not None:
        updates.append("is_active = ?")
        params.append(1 if req.is_active else 0)
    if updates:
        params.append(announcement_id)
        conn.execute(f"UPDATE announcements SET {', '.join(updates)} WHERE id = ?", params)
        conn.commit()
    conn.close()
    return {"ok": True}

@app.delete("/api/announcements/{announcement_id}")
async def delete_announcement(announcement_id: str, current_user: dict = Depends(get_admin_user)):
    conn = get_db()
    conn.execute("DELETE FROM announcement_reads WHERE announcement_id = ?", (announcement_id,))
    conn.execute("DELETE FROM announcements WHERE id = ?", (announcement_id,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.post("/api/announcements/{announcement_id}/read")
async def mark_announcement_read(announcement_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id FROM announcements WHERE id = ?", (announcement_id,))
    if not c.fetchone():
        conn.close()
        raise HTTPException(404, "鍏憡涓嶅瓨鍦?")
    now = datetime.now().isoformat()
    conn.execute(
        "INSERT OR REPLACE INTO announcement_reads (announcement_id, username, read_at) VALUES (?, ?, ?)",
        (announcement_id, current_user["username"], now)
    )
    conn.commit()
    conn.close()
    return {"ok": True}

# ============ Workspace API ============

@app.get("/api/workspaces")
async def list_workspaces(current_user: dict = Depends(get_user)):
    conn = get_db()
    # Always return ALL workspaces; attach membership info per user
    rows = conn.execute(
        "SELECT w.*, u.name as creator_name FROM workspaces w LEFT JOIN users u ON w.created_by = u.username ORDER BY w.created_at DESC"
    ).fetchall()
    uname = current_user["username"]
    is_admin = current_user["role"] == "admin"
    result = []
    for r in rows:
        ws = dict(r)
        ws["member_count"] = conn.execute(
            "SELECT COUNT(*) as cnt FROM workspace_members WHERE workspace_id = ?", (r["id"],)
        ).fetchone()["cnt"]
        my_row = conn.execute(
            "SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (r["id"], uname)
        ).fetchone()
        if my_row:
            ws["is_member"] = True
            ws["my_role"] = my_row["role"]
        elif is_admin:
            ws["is_member"] = True
            ws["my_role"] = "admin"
        else:
            ws["is_member"] = False
            ws["my_role"] = None
        # Check for pending join request
        req_row = conn.execute(
            "SELECT id FROM workspace_join_requests WHERE workspace_id = ? AND username = ? AND status = 'pending'",
            (r["id"], uname)
        ).fetchone()
        ws["has_pending_request"] = bool(req_row)
        result.append(ws)
    conn.close()
    return result

@app.post("/api/workspaces")
async def create_workspace(req: WorkspaceCreate, current_user: dict = Depends(get_admin_user)):
    conn = get_db()
    if conn.execute("SELECT 1 FROM workspaces WHERE name = ?", (req.name,)).fetchone():
        conn.close()
        raise HTTPException(400, "宸ヤ綔绌洪棿鍚嶇О宸插瓨鍦?")
    wid = uuid.uuid4().hex[:12]
    now = datetime.now().isoformat()
    conn.execute(
        "INSERT INTO workspaces (id, name, description, created_at, created_by, quota_limit) VALUES (?, ?, ?, ?, ?, ?)",
        (wid, req.name, req.description, now, current_user["username"], req.quota_limit)
    )
    conn.execute(
        "INSERT OR IGNORE INTO workspace_members (workspace_id, username, role, joined_at, quota_limit) VALUES (?, ?, 'admin', ?, ?)",
        (wid, current_user["username"], now, WORKSPACE_QUOTA_DEFAULT) # Workspace admin gets full quota by default? 
        # Actually the prompt says: 宸ヤ綔绌洪棿鐨勭鐞嗗憳鏄嫢鏈?00G鐨勬敮閰嶆潈鍒嗛厤缁欎笅灞炵殑鏅€氭垚鍛?
        # So maybe admins should have WORKSPACE_QUOTA_DEFAULT as their limit too?
    )
    conn.commit()
    conn.close()
    return {"ok": True, "id": wid}

@app.patch("/api/workspaces/{workspace_id}")
async def update_workspace(workspace_id: str, req: WorkspaceUpdate, current_user: dict = Depends(get_admin_user)):
    conn = get_db()
    if not conn.execute("SELECT 1 FROM workspaces WHERE id = ?", (workspace_id,)).fetchone():
        conn.close()
        raise HTTPException(404, "宸ヤ綔绌洪棿涓嶅瓨鍦?")
    updates, params = [], []
    if req.name is not None:
        updates.append("name = ?"); params.append(req.name)
    if req.description is not None:
        updates.append("description = ?"); params.append(req.description)
    if updates:
        params.append(workspace_id)
        conn.execute(f"UPDATE workspaces SET {', '.join(updates)} WHERE id = ?", params)
        conn.commit()
    conn.close()
    return {"ok": True}

@app.delete("/api/workspaces/{workspace_id}")
async def delete_workspace(workspace_id: str, current_user: dict = Depends(get_admin_user)):
    conn = get_db()
    if not conn.execute("SELECT 1 FROM workspaces WHERE id = ?", (workspace_id,)).fetchone():
        conn.close()
        raise HTTPException(404, "宸ヤ綔绌洪棿涓嶅瓨鍦?")
    file_count = conn.execute(
        "SELECT COUNT(*) as cnt FROM files WHERE workspace_id = ?", (workspace_id,)
    ).fetchone()["cnt"]
    if file_count > 0:
        conn.close()
        raise HTTPException(400, f"宸ヤ綔绌洪棿鍐呰繕鏈?{file_count} 涓枃浠讹紝璇峰厛绉婚櫎鎴栬浆绉绘枃浠?")
    conn.execute("DELETE FROM workspace_members WHERE workspace_id = ?", (workspace_id,))
    conn.execute("DELETE FROM workspaces WHERE id = ?", (workspace_id,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/api/workspaces/{workspace_id}/members")
async def list_workspace_members(workspace_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    if current_user["role"] != "admin":
        if not conn.execute(
            "SELECT 1 FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (workspace_id, current_user["username"])
        ).fetchone():
            conn.close()
            raise HTTPException(403, "鏃犳潈闄愭煡鐪嬫宸ヤ綔绌洪棿")
    rows = conn.execute(
        """SELECT wm.username, wm.role, wm.joined_at, u.name
           FROM workspace_members wm JOIN users u ON wm.username = u.username
           WHERE wm.workspace_id = ?
           ORDER BY wm.role DESC, wm.joined_at ASC""",
        (workspace_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/api/workspaces/{workspace_id}/members")
async def add_workspace_member(workspace_id: str, req: WorkspaceMemberUpdate, current_user: dict = Depends(get_user)):
    conn = get_db()
    if current_user["role"] != "admin":
        ws_role = conn.execute(
            "SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (workspace_id, current_user["username"])
        ).fetchone()
        if not ws_role or ws_role["role"] != "admin":
            conn.close()
            raise HTTPException(403, "闇€瑕佸伐浣滅┖闂寸鐞嗗憳鏉冮檺")
    if not conn.execute("SELECT 1 FROM workspaces WHERE id = ?", (workspace_id,)).fetchone():
        conn.close()
        raise HTTPException(404, "宸ヤ綔绌洪棿涓嶅瓨鍦?")
    if not conn.execute("SELECT 1 FROM users WHERE username = ? AND status = 'active'", (req.username,)).fetchone():
        conn.close()
        raise HTTPException(404, "鐢ㄦ埛涓嶅瓨鍦ㄦ垨鏈縺娲?")
    if req.role not in ("admin", "member"):
        conn.close()
        raise HTTPException(400, "瑙掕壊鏃犳晥")
    now = datetime.now().isoformat()
    conn.execute(
        "INSERT OR REPLACE INTO workspace_members (workspace_id, username, role, joined_at) VALUES (?, ?, ?, ?)",
        (workspace_id, req.username, req.role, now)
    )
    conn.commit()
    conn.close()
    return {"ok": True}

@app.delete("/api/workspaces/{workspace_id}/members/{username}")
async def remove_workspace_member(workspace_id: str, username: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    if current_user["role"] != "admin":
        ws_role = conn.execute(
            "SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (workspace_id, current_user["username"])
        ).fetchone()
        if not ws_role or ws_role["role"] != "admin":
            conn.close()
            raise HTTPException(403, "闇€瑕佸伐浣滅┖闂寸鐞嗗憳鏉冮檺")
    conn.execute(
        "DELETE FROM workspace_members WHERE workspace_id = ? AND username = ?", (workspace_id, username)
    )
    conn.commit()
    conn.close()
    return {"ok": True}

@app.patch("/api/workspaces/{workspace_id}/members/{username}")
async def update_workspace_member_role(workspace_id: str, username: str, req: dict, current_user: dict = Depends(get_admin_user)):
    role = req.get("role")
    if role not in ("admin", "member"):
        raise HTTPException(400, "瑙掕壊鏃犳晥")
    conn = get_db()
    conn.execute(
        "UPDATE workspace_members SET role = ? WHERE workspace_id = ? AND username = ?",
        (role, workspace_id, username)
    )
    conn.commit()
    conn.close()
    return {"ok": True}

# ============ Workspace Join Requests API ============

@app.post("/api/workspaces/{workspace_id}/request-join")
async def request_join_workspace(workspace_id: str, current_user: dict = Depends(get_user)):
    if current_user["role"] == "admin":
        raise HTTPException(400, "绠＄悊鍛樻棤闇€鐢宠")
    conn = get_db()
    if not conn.execute("SELECT 1 FROM workspaces WHERE id = ?", (workspace_id,)).fetchone():
        conn.close()
        raise HTTPException(404, "宸ヤ綔绌洪棿涓嶅瓨鍦?")
    if conn.execute(
        "SELECT 1 FROM workspace_members WHERE workspace_id = ? AND username = ?",
        (workspace_id, current_user["username"])
    ).fetchone():
        conn.close()
        raise HTTPException(400, "鎮ㄥ凡鏄宸ヤ綔绌洪棿鎴愬憳")
    rid = uuid.uuid4().hex[:16]
    now = datetime.now().isoformat()
    try:
        conn.execute(
            "INSERT INTO workspace_join_requests (id, workspace_id, username, status, requested_at) VALUES (?, ?, ?, 'pending', ?)",
            (rid, workspace_id, current_user["username"], now)
        )
    except Exception:
        # UNIQUE constraint: already has a request, update to pending
        conn.execute(
            "UPDATE workspace_join_requests SET status='pending', requested_at=?, reviewed_by='', reviewed_at='' WHERE workspace_id=? AND username=?",
            (now, workspace_id, current_user["username"])
        )
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/api/todos")
async def get_todos(current_user: dict = Depends(get_user)):
    """Returns pending todos: user approvals (admin only) + workspace join requests (admin or ws-admin) + file transfers (all users)."""
    conn = get_db()
    result = {"user_approvals": [], "ws_join_requests": [], "file_transfers": []}

    if current_user["role"] == "admin":
        # Pending user registrations
        rows = conn.execute(
            "SELECT username, name, role, status, created_at FROM users WHERE status = 'pending' ORDER BY created_at DESC"
        ).fetchall()
        result["user_approvals"] = [dict(r) for r in rows]
        # All pending workspace join requests
        rows = conn.execute("""
            SELECT r.id, r.workspace_id, r.username, r.requested_at,
                   w.name as workspace_name, u.name as user_display_name
            FROM workspace_join_requests r
            JOIN workspaces w ON w.id = r.workspace_id
            JOIN users u ON u.username = r.username
            WHERE r.status = 'pending'
            ORDER BY r.requested_at DESC
        """).fetchall()
        result["ws_join_requests"] = [dict(r) for r in rows]
    else:
        # Workspace admins: only see requests for their workspaces
        admin_ws = conn.execute(
            "SELECT workspace_id FROM workspace_members WHERE username = ? AND role = 'admin'",
            (current_user["username"],)
        ).fetchall()
        if admin_ws:
            ids = [r["workspace_id"] for r in admin_ws]
            ph = ",".join("?" * len(ids))
            rows = conn.execute(f"""
                SELECT r.id, r.workspace_id, r.username, r.requested_at,
                       w.name as workspace_name, u.name as user_display_name
                FROM workspace_join_requests r
                JOIN workspaces w ON w.id = r.workspace_id
                JOIN users u ON u.username = r.username
                WHERE r.status = 'pending' AND r.workspace_id IN ({ph})
                ORDER BY r.requested_at DESC
            """, ids).fetchall()
            result["ws_join_requests"] = [dict(r) for r in rows]

    # File transfers: all users see their incoming pending transfers
    transfers = conn.execute("""
        SELECT ft.id, ft.file_id, ft.sender, ft.message, ft.created_at,
               f.filename, f.type as file_type, u.name as sender_name
        FROM file_transfers ft
        JOIN files f ON f.id = ft.file_id
        JOIN users u ON u.username = ft.sender
        WHERE ft.recipient = ? AND ft.status = 'pending'
        ORDER BY ft.created_at DESC
    """, (current_user["username"],)).fetchall()
    result["file_transfers"] = [dict(r) for r in transfers]

    conn.close()
    return result

@app.post("/api/todos/ws-requests/{request_id}/approve")
async def approve_ws_request(request_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    req = conn.execute(
        "SELECT * FROM workspace_join_requests WHERE id = ?", (request_id,)
    ).fetchone()
    if not req or req["status"] != "pending":
        conn.close()
        raise HTTPException(404, "鐢宠涓嶅瓨鍦ㄦ垨宸插鐞?")
    # Check permission
    if current_user["role"] != "admin":
        is_ws_admin = conn.execute(
            "SELECT 1 FROM workspace_members WHERE workspace_id = ? AND username = ? AND role = 'admin'",
            (req["workspace_id"], current_user["username"])
        ).fetchone()
        if not is_ws_admin:
            conn.close()
            raise HTTPException(403, "娌℃湁鏉冮檺")
    now = datetime.now().isoformat()
    conn.execute(
        "UPDATE workspace_join_requests SET status='approved', reviewed_by=?, reviewed_at=? WHERE id=?",
        (current_user["username"], now, request_id)
    )
    # Add to workspace_members
    if not conn.execute(
        "SELECT 1 FROM workspace_members WHERE workspace_id=? AND username=?",
        (req["workspace_id"], req["username"])
    ).fetchone():
        conn.execute(
            "INSERT INTO workspace_members (workspace_id, username, role, joined_at) VALUES (?, ?, 'member', ?)",
            (req["workspace_id"], req["username"], now)
        )
    conn.commit()
    conn.close()
    log_admin_action(current_user["username"], "approve_ws_join", req["username"], f"workspace: {req['workspace_id']}")
    return {"ok": True}

@app.post("/api/todos/ws-requests/{request_id}/reject")
async def reject_ws_request(request_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    req = conn.execute(
        "SELECT * FROM workspace_join_requests WHERE id = ?", (request_id,)
    ).fetchone()
    if not req or req["status"] != "pending":
        conn.close()
        raise HTTPException(404, "鐢宠涓嶅瓨鍦ㄦ垨宸插鐞?")
    if current_user["role"] != "admin":
        is_ws_admin = conn.execute(
            "SELECT 1 FROM workspace_members WHERE workspace_id = ? AND username = ? AND role = 'admin'",
            (req["workspace_id"], current_user["username"])
        ).fetchone()
        if not is_ws_admin:
            conn.close()
            raise HTTPException(403, "娌℃湁鏉冮檺")
    now = datetime.now().isoformat()
    conn.execute(
        "UPDATE workspace_join_requests SET status='rejected', reviewed_by=?, reviewed_at=? WHERE id=?",
        (current_user["username"], now, request_id)
    )
    conn.commit()
    conn.close()
    log_admin_action(current_user["username"], "reject_ws_join", req["username"], f"workspace: {req['workspace_id']}")
    return {"ok": True}

# ============ File Transfer API ============

@app.get("/api/users/active-list")
async def list_active_users(current_user: dict = Depends(get_user)):
    """Return basic info of all active users. Available to all authenticated users for file transfer."""
    conn = get_db()
    rows = conn.execute(
        "SELECT username, name FROM users WHERE status = 'active' ORDER BY name"
    ).fetchall()
    conn.close()
    return [{"username": r["username"], "name": r["name"]} for r in rows]

@app.post("/api/files/{file_id}/transfer")
async def api_transfer_file(file_id: str, req: FileTransferRequest, current_user: dict = Depends(get_user)):
    """Send a file to another user."""
    conn = get_db()
    file_row = conn.execute("SELECT id, filename FROM files WHERE id = ?", (file_id,)).fetchone()
    if not file_row:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    if req.recipient == current_user["username"]:
        conn.close()
        raise HTTPException(400, "涓嶈兘鍙戦€佺粰鑷繁")
    recipient_row = conn.execute(
        "SELECT username FROM users WHERE username = ? AND status = 'active'", (req.recipient,)
    ).fetchone()
    if not recipient_row:
        conn.close()
        raise HTTPException(404, "鐢ㄦ埛涓嶅瓨鍦ㄦ垨鏈縺娲?")
    transfer_id = uuid.uuid4().hex[:16]
    now = datetime.now().isoformat()
    conn.execute(
        "INSERT INTO file_transfers (id, file_id, sender, recipient, message, status, created_at) VALUES (?, ?, ?, ?, ?, 'pending', ?)",
        (transfer_id, file_id, current_user["username"], req.recipient, req.message, now)
    )
    conn.commit()
    conn.close()
    return {"ok": True, "transfer_id": transfer_id}

@app.post("/api/todos/file-transfers/{transfer_id}/accept")
async def api_accept_file_transfer(transfer_id: str, req: FileTransferAcceptRequest, current_user: dict = Depends(get_user)):
    """Accept a file transfer, copying the file into selected workspace(s)."""
    conn = get_db()
    transfer = conn.execute(
        "SELECT * FROM file_transfers WHERE id = ? AND recipient = ? AND status = 'pending'",
        (transfer_id, current_user["username"])
    ).fetchone()
    if not transfer:
        conn.close()
        raise HTTPException(404, "浼犻€掔敵璇蜂笉瀛樺湪鎴栧凡澶勭悊")
    orig = conn.execute("SELECT * FROM files WHERE id = ?", (transfer["file_id"],)).fetchone()
    if not orig:
        conn.close()
        raise HTTPException(404, "鍘熷鏂囦欢宸蹭笉瀛樺湪")

    now = datetime.now().isoformat()
    target_ws_ids = req.workspace_ids if req.workspace_ids else [""]
    first_new_id = None

    for ws_id in target_ws_ids:
        new_id = uuid.uuid4().hex[:12]
        ext = orig["type"]
        src_path = UPLOAD_DIR / f"{orig['id']}{ext}"
        dst_path = UPLOAD_DIR / f"{new_id}{ext}"
        if src_path.exists():
            shutil.copy2(str(src_path), str(dst_path))
            new_size = dst_path.stat().st_size
        else:
            new_size = orig["size"]
        conn.execute(
            "INSERT INTO files (id, filename, size, type, folder, workspace_id, view_count, uploaded_at, uploaded_by, analysis) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (new_id, orig["filename"], new_size, ext, "", ws_id, 0, now, current_user["username"], orig["analysis"])
        )
        _write_file_perms(new_id, req.visibility, req.allow_view, req.allow_edit, req.allow_delete, conn)
        conn.execute(
            "INSERT OR IGNORE INTO file_favorites (username, file_id, created_at) VALUES (?, ?, ?)",
            (current_user["username"], new_id, now)
        )
        if first_new_id is None:
            first_new_id = new_id

    conn.execute(
        "UPDATE file_transfers SET status='accepted', resolved_at=? WHERE id=?",
        (now, transfer_id)
    )
    conn.commit()
    conn.close()
    return {"ok": True, "file_id": first_new_id}

@app.post("/api/todos/file-transfers/{transfer_id}/reject")
async def api_reject_file_transfer(transfer_id: str, current_user: dict = Depends(get_user)):
    """Reject a file transfer."""
    conn = get_db()
    transfer = conn.execute(
        "SELECT * FROM file_transfers WHERE id = ? AND recipient = ? AND status = 'pending'",
        (transfer_id, current_user["username"])
    ).fetchone()
    if not transfer:
        conn.close()
        raise HTTPException(404, "浼犻€掔敵璇蜂笉瀛樺湪鎴栧凡澶勭悊")
    now = datetime.now().isoformat()
    conn.execute(
        "UPDATE file_transfers SET status='rejected', resolved_at=? WHERE id=?",
        (now, transfer_id)
    )
    conn.commit()
    conn.close()
    return {"ok": True}

# ============ File Permissions API ============

@app.get("/api/files/{file_id}/permissions")
async def get_file_permissions_api(file_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    f = conn.execute("SELECT uploaded_by, workspace_id FROM files WHERE id = ?", (file_id,)).fetchone()
    if not f:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    if current_user["role"] != "admin" and f["uploaded_by"] != current_user["username"]:
        ws_role = conn.execute(
            "SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (f["workspace_id"], current_user["username"])
        ).fetchone()
        if not ws_role or ws_role["role"] != "admin":
            conn.close()
            raise HTTPException(403, "鏃犳潈闄?")
    perms = get_file_perms(file_id, conn)
    conn.close()
    return perms

@app.put("/api/files/{file_id}/permissions")
async def update_file_permissions_api(file_id: str, req: FilePermissionsUpdate, current_user: dict = Depends(get_user)):
    conn = get_db()
    f = conn.execute("SELECT uploaded_by, workspace_id FROM files WHERE id = ?", (file_id,)).fetchone()
    if not f:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    if current_user["role"] != "admin" and f["uploaded_by"] != current_user["username"]:
        ws_role = conn.execute(
            "SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (f["workspace_id"], current_user["username"])
        ).fetchone()
        if not ws_role or ws_role["role"] != "admin":
            conn.close()
            raise HTTPException(403, "鏃犳潈闄愪慨鏀规枃浠舵潈闄?")
    if req.visibility not in ("private", "workspace", "public"):
        conn.close()
        raise HTTPException(400, "visibility 鏃犳晥")
    _write_file_perms(file_id, req.visibility, req.allow_view, req.allow_edit, req.allow_delete, conn)
    conn.commit()
    conn.close()
    return {"ok": True}

# User management (admin only)
class PermissionUpdateRequest(BaseModel):
    can_upload: bool = True
    can_delete_own: bool = True
    can_delete_any: bool = False
    can_edit_own: bool = True
    can_edit_others: bool = False
    can_create_folder: bool = False
    can_view_report: bool = False

@app.get("/api/users")
async def list_users(q: Optional[str] = None, workspace_id: Optional[str] = None, authorization: Optional[str] = Header(None)):
    user = get_user(authorization)
    conn = get_db()
    c = conn.cursor()
    is_sys_admin = user["role"] == "admin"

    # Determine which workspace_ids the caller is allowed to filter
    if not is_sys_admin:
        c.execute("SELECT workspace_id FROM workspace_members WHERE username = ? AND role = 'admin'", (user["username"],))
        admin_ws_ids = [r["workspace_id"] for r in c.fetchall()]
        if not admin_ws_ids:
            conn.close()
            raise HTTPException(403, "娌℃湁鏉冮檺")
        # If specific workspace requested, verify caller is admin of it
        if workspace_id and workspace_id not in admin_ws_ids:
            conn.close()
            raise HTTPException(403, "娌℃湁鏉冮檺璁块棶璇ュ伐浣滅┖闂?")
        # If no filter, restrict to union of their admin workspaces
        if not workspace_id:
            workspace_id = "__ws_admin_filter__"
            filter_ws_ids = admin_ws_ids
        else:
            filter_ws_ids = [workspace_id]
    else:
        filter_ws_ids = [workspace_id] if workspace_id else None

    like = f"%{q}%" if q else None

    if filter_ws_ids is not None:
        ph = ",".join("?" * len(filter_ws_ids))
        base_args = filter_ws_ids
        if like:
            rows = conn.execute(f"""
                SELECT DISTINCT u.username, u.name, u.role, u.status, u.created_at, u.bio
                FROM users u
                INNER JOIN workspace_members wm ON u.username = wm.username AND wm.workspace_id IN ({ph})
                WHERE (u.username LIKE ? OR u.name LIKE ?)
            """, base_args + [like, like]).fetchall()
        else:
            rows = conn.execute(f"""
                SELECT DISTINCT u.username, u.name, u.role, u.status, u.created_at, u.bio
                FROM users u
                INNER JOIN workspace_members wm ON u.username = wm.username AND wm.workspace_id IN ({ph})
            """, base_args).fetchall()
    else:
        if like:
            rows = conn.execute("SELECT username, name, role, status, created_at, bio FROM users WHERE username LIKE ? OR name LIKE ?", (like, like)).fetchall()
        else:
            rows = conn.execute("SELECT username, name, role, status, created_at, bio FROM users").fetchall()

    # Attach workspace memberships for each returned user
    usernames = [r["username"] for r in rows]
    ws_map: dict = {}
    if usernames:
        ph2 = ",".join("?" * len(usernames))
        ws_rows = conn.execute(f"""
            SELECT wm.username, wm.role AS ws_role, w.id AS ws_id, w.name AS ws_name
            FROM workspace_members wm
            JOIN workspaces w ON w.id = wm.workspace_id
            WHERE wm.username IN ({ph2})
        """, usernames).fetchall()
        for wr in ws_rows:
            ws_map.setdefault(wr["username"], []).append({
                "id": wr["ws_id"], "name": wr["ws_name"], "role": wr["ws_role"]
            })

    conn.close()
    return [{
        "username": r["username"], "name": r["name"], "role": r["role"],
        "status": r["status"], "created_at": r["created_at"], "bio": r["bio"] or "",
        "permissions": get_permissions(r["username"]),
        "workspaces": ws_map.get(r["username"], []),
    } for r in rows]

@app.post("/api/users")
async def create_user(req: dict, authorization: Optional[str] = Header(None)):
    actor = get_admin_user(authorization)
    username = req.get("username")
    password = req.get("password")
    name = req.get("name")
    role = req.get("role", "member")
    skip_weak_check = req.get("skip_weak_check", False)

    if not username or not password:
        raise HTTPException(400, "鐢ㄦ埛鍚嶅拰瀵嗙爜涓嶈兘涓虹┖")

    # Admin privilege: warning but allowed to skip
    if not is_strong_password(password) and not skip_weak_check:
        return JSONResponse(status_code=400, content={"detail": "WEAK_PASSWORD_WARNING", "message": "璇ュ瘑鐮佸己搴﹁緝寮憋紙寤鸿鍖呭惈瀛楁瘝鍜屾暟瀛椾笖涓嶅皯浜?浣嶏級銆傛槸鍚﹀潥鎸佷娇鐢ㄦ瀵嗙爜锛?"})

    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) as cnt FROM users WHERE username = ?", (username,))
    if c.fetchone()["cnt"] > 0:
        conn.close()
        raise HTTPException(400, "鐢ㄦ埛鍚嶅凡瀛樺湪")
    
    c.execute("INSERT INTO users (username, password_hash, name, role, status, created_at) VALUES (?, ?, ?, ?, ?, ?)",
              (username, hash_pw(password), name or username, role, "active", datetime.now().isoformat()))
    conn.commit()
    conn.close()
    log_admin_action(actor["username"], "create_user", username, f"Role: {role}")
    return {"ok": True}

@app.delete("/api/users/{username}")
async def delete_user(username: str, authorization: Optional[str] = Header(None)):
    actor = get_admin_user(authorization)
    if actor["username"] == username:
        raise HTTPException(400, "涓嶈兘鍒犻櫎鑷繁")
    conn = get_db()
    conn.execute("DELETE FROM users WHERE username = ?", (username,))
    conn.commit()
    conn.close()
    log_admin_action(actor["username"], "delete_user", username)
    return {"ok": True}

@app.post("/api/users/{username}/role")
async def update_user_role(username: str, role_data: dict, authorization: Optional[str] = Header(None)):
    actor = get_admin_user(authorization)
    if role_data.get("role") not in ("admin", "member"):
        raise HTTPException(400, "瑙掕壊鏃犳晥")
    conn = get_db()
    conn.execute("UPDATE users SET role = ? WHERE username = ?", (role_data["role"], username))
    conn.commit()
    conn.close()
    log_admin_action(actor["username"], "update_role", username, f"New Role: {role_data['role']}")
    return {"ok": True}

@app.get("/api/users/{username}/permissions")
async def get_user_permissions(username: str, authorization: Optional[str] = Header(None)):
    get_admin_user(authorization)
    perms = get_permissions(username)
    return perms

@app.put("/api/users/{username}/permissions")
async def update_user_permissions(username: str, req: PermissionUpdateRequest, authorization: Optional[str] = Header(None)):
    actor = get_admin_user(authorization)
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT role FROM users WHERE username = ?", (username,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鐢ㄦ埛涓嶅瓨鍦?")
    if row["role"] == "admin":
        conn.close()
        raise HTTPException(400, "绠＄悊鍛樻潈闄愪笉鍙慨鏀?")
    perms = {
        "can_upload": req.can_upload,
        "can_delete_own": req.can_delete_own,
        "can_delete_any": req.can_delete_any,
        "can_edit_own": req.can_edit_own,
        "can_edit_others": req.can_edit_others,
        "can_create_folder": req.can_create_folder,
        "can_view_report": req.can_view_report,
    }
    perms_json = json.dumps(perms, ensure_ascii=False)
    c.execute("UPDATE users SET permissions = ? WHERE username = ?", (perms_json, username))
    conn.commit()
    conn.close()
    log_admin_action(actor["username"], "update_permissions", username, perms_json)
    return {"ok": True, "permissions": perms}

@app.get("/api/users/pending")
async def list_pending_users(authorization: Optional[str] = Header(None)):
    get_admin_user(authorization)
    conn = get_db()
    rows = conn.execute("SELECT username, name, role, status, created_at FROM users WHERE status = 'pending'").fetchall()
    conn.close()
    return [{"username": r["username"], "name": r["name"], "role": r["role"], "status": "pending", "created_at": r["created_at"]} for r in rows]

@app.post("/api/users/{username}/approve")
async def approve_user(username: str, authorization: Optional[str] = Header(None)):
    actor = get_admin_user(authorization)
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT status FROM users WHERE username = ?", (username,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鐢ㄦ埛涓嶅瓨鍦?")
    if row["status"] == "active":
        conn.close()
        raise HTTPException(400, "鐢ㄦ埛宸叉縺娲?")
    conn.execute("UPDATE users SET status = 'active' WHERE username = ?", (username,))
    conn.commit()
    conn.close()
    log_admin_action(actor["username"], "approve_user", username)
    return {"ok": True}

@app.post("/api/users/{username}/reject")
async def reject_user(username: str, authorization: Optional[str] = Header(None)):
    actor = get_admin_user(authorization)
    if actor["username"] == username:
        raise HTTPException(400, "涓嶈兘鎷掔粷鑷繁")
    conn = get_db()
    conn.execute("DELETE FROM users WHERE username = ?", (username,))
    conn.commit()
    conn.close()
    log_admin_action(actor["username"], "reject_user", username)
    return {"ok": True}

# ============ Workspace Access Helpers ============

def get_file_perms(file_id: str, conn) -> dict:
    row = conn.execute("SELECT * FROM file_permissions WHERE file_id = ?", (file_id,)).fetchone()
    if not row:
        return {"visibility": "workspace", "allow_view": True, "allow_edit": False, "allow_delete": False}
    return {
        "visibility": row["visibility"],
        "allow_view": bool(row["allow_view"]),
        "allow_edit": bool(row["allow_edit"]),
        "allow_delete": bool(row["allow_delete"]),
    }

def resolve_file_access(file_id: str, workspace_id: str, uploaded_by: str, current_user: dict, conn) -> dict:
    """Compute visibility and permission flags for a file given the current user."""
    username = current_user["username"]

    if current_user["role"] == "admin":
        return {"visible": True, "can_view": True, "can_edit": True, "can_delete": True}

    if uploaded_by == username:
        return {"visible": True, "can_view": True, "can_edit": True, "can_delete": True}

    if workspace_id:
        ws_role = conn.execute(
            "SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (workspace_id, username)
        ).fetchone()
        if ws_role and ws_role["role"] == "admin":
            return {"visible": True, "can_view": True, "can_edit": True, "can_delete": True}

    perms = get_file_perms(file_id, conn)

    if not workspace_id:
        return {
            "visible": True,
            "can_view": perms["allow_view"],
            "can_edit": perms["allow_edit"],
            "can_delete": perms["allow_delete"],
        }

    if perms["visibility"] == "private":
        return {"visible": False, "can_view": False, "can_edit": False, "can_delete": False}

    if perms["visibility"] == "workspace":
        in_ws = conn.execute(
            "SELECT 1 FROM workspace_members WHERE workspace_id = ? AND username = ?",
            (workspace_id, username)
        ).fetchone()
        if not in_ws:
            return {"visible": False, "can_view": False, "can_edit": False, "can_delete": False}

    return {
        "visible": True,
        "can_view": perms["allow_view"],
        "can_edit": perms["allow_edit"],
        "can_delete": perms["allow_delete"],
    }

def _write_file_perms(file_id: str, visibility: str, allow_view: bool, allow_edit: bool, allow_delete: bool, conn):
    conn.execute(
        "INSERT OR REPLACE INTO file_permissions (file_id, visibility, allow_view, allow_edit, allow_delete) VALUES (?, ?, ?, ?, ?)",
        (file_id, visibility, 1 if allow_view else 0, 1 if allow_edit else 0, 1 if allow_delete else 0)
    )

# ============ File Store ============

def file_to_dict(row):
    """Convert sqlite Row to dict matching old JSON structure."""
    keys = row.keys() if hasattr(row, "keys") else []
    return {
        "id": row["id"],
        "filename": row["filename"],
        "size": row["size"],
        "type": row["type"],
        "folder": row["folder"],
        "workspace_id": row["workspace_id"] if "workspace_id" in keys else "",
        "view_count": row["view_count"],
        "uploaded_at": row["uploaded_at"],
        "uploaded_by": row["uploaded_by"],
        "last_modified_at": row["last_modified_at"] if "last_modified_at" in keys else None,
        "last_modified_by": row["last_modified_by"] if "last_modified_by" in keys else None,
        "analysis": json.loads(row["analysis"]) if row["analysis"] else {},
    }

def graph_node_to_dict(row) -> dict:
    item = dict(row)
    try:
        analysis = json.loads(row["analysis"]) if row["analysis"] else {}
    except Exception:
        analysis = {}
    content_amount = 0
    if isinstance(analysis, dict):
        content_amount = (
            analysis.get("char_count")
            or analysis.get("word_count")
            or ((analysis.get("rows") or 0) * max(1, analysis.get("columns") or 1))
            or analysis.get("line_count")
            or 0
        )
    item["content_amount"] = int(content_amount or 0)
    item["analysis_summary"] = {
        "char_count": analysis.get("char_count", 0) if isinstance(analysis, dict) else 0,
        "word_count": analysis.get("word_count", 0) if isinstance(analysis, dict) else 0,
        "line_count": analysis.get("line_count", 0) if isinstance(analysis, dict) else 0,
        "rows": analysis.get("rows", 0) if isinstance(analysis, dict) else 0,
        "columns": analysis.get("columns", 0) if isinstance(analysis, dict) else 0,
    }
    return item

def normalize_file_link_pair(file_id: str, related_file_id: str) -> tuple[str, str]:
    return tuple(sorted((file_id, related_file_id)))

def can_user_access_file_row(row, current_user: dict, conn) -> bool:
    ws = row["workspace_id"] if "workspace_id" in row.keys() else ""
    access = resolve_file_access(row["id"], ws, row["uploaded_by"], current_user, conn)
    return bool(access["visible"] and access["can_view"])

def write_file_link_log(link_id: str, action: str, file_id: str, related_file_id: str, relation_label: str, created_by: str, conn):
    conn.execute(
        """
        INSERT INTO file_link_logs
        (id, link_id, action, file_id, related_file_id, relation_label, created_at, created_by)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            uuid.uuid4().hex[:16],
            link_id,
            action,
            file_id,
            related_file_id,
            relation_label or "",
            datetime.now().isoformat(),
            created_by,
        )
    )

def save_file_version(file_id: str, saved_by: str):
    """Save current file as a version before updating."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT filename, size, type, workspace_id, uploaded_by FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        return

    ext = row["type"]
    current_file = UPLOAD_DIR / f"{file_id}{ext}"
    if not current_file.exists():
        conn.close()
        return

    # Generate version info
    c.execute("SELECT MAX(version_num) as v FROM file_versions WHERE file_id = ?", (file_id,))
    last_v = c.fetchone()["v"] or 0
    new_v = last_v + 1

    version_id = uuid.uuid4().hex[:16]
    version_file = UPLOAD_VERSIONS_DIR / f"{version_id}{ext}"
    shutil.copy2(current_file, version_file)

    version_size = current_file.stat().st_size
    c.execute("INSERT INTO file_versions (id, file_id, version_num, size, saved_at, saved_by) VALUES (?, ?, ?, ?, ?, ?)",
              (version_id, file_id, new_v, version_size, datetime.now().isoformat(), saved_by))

    # Update quota for the new version
    update_quota(row["uploaded_by"], row["workspace_id"], version_size, conn)

    conn.commit()
    conn.close()
@app.get("/api/files/{file_id}/versions")
async def list_versions(file_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    rows = conn.execute("SELECT * FROM file_versions WHERE file_id = ? ORDER BY version_num DESC", (file_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/api/files/{file_id}/versions/{version_id}/download")
async def download_version(file_id: str, version_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    v = conn.execute("SELECT * FROM file_versions WHERE id = ? AND file_id = ?", (version_id, file_id)).fetchone()
    if not v:
        conn.close()
        raise HTTPException(404, "鐗堟湰涓嶅瓨鍦?")
    
    f = conn.execute("SELECT type, filename FROM files WHERE id = ?", (file_id,)).fetchone()
    conn.close()
    
    ext = f["type"].lower()
    version_file = UPLOAD_VERSIONS_DIR / f"{version_id}{ext}"
    if not version_file.exists():
        raise HTTPException(404, "鐗堟湰鏂囦欢涓㈠け")
        
    return FileResponse(
        path=version_file,
        filename=f"V{v['version_num']}_{f['filename']}",
        media_type="application/octet-stream",
    )

@app.post("/api/files/{file_id}/versions/{version_id}/restore")
async def restore_version(file_id: str, version_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    v = conn.execute("SELECT * FROM file_versions WHERE id = ? AND file_id = ?", (version_id, file_id)).fetchone()
    if not v:
        conn.close()
        raise HTTPException(404, "鐗堟湰涓嶅瓨鍦?")
    
    f = conn.execute("SELECT type, filename FROM files WHERE id = ?", (file_id,)).fetchone()
    conn.close()
    
    ext = f["type"].lower()
    version_file = UPLOAD_VERSIONS_DIR / f"{version_id}{ext}"
    current_file = UPLOAD_DIR / f"{file_id}{ext}"
    
    if not version_file.exists():
        raise HTTPException(404, "鐗堟湰鏂囦欢涓㈠け")
        
    # Before restoring, save CURRENT state as a NEW version
    save_file_version(file_id, current_user["username"])
    
    # Copy version file to current file
    shutil.copy2(version_file, current_file)
    
    # Update current file metadata
    new_size = current_file.stat().st_size
    size_change = new_size - v["size"] # Wait, v["size"] is the size of the version we are restoring.
    # Actually, old_size is what we replaced.
    # The flow is: 
    # 1. save_file_version(current_file) -> quota increases by current_file.size
    # 2. current_file is replaced by version_file
    # 3. size_change = version_file.size - current_file.size
    
    old_size = conn.execute("SELECT size FROM files WHERE id = ?", (file_id,)).fetchone()["size"]
    new_size = current_file.stat().st_size
    size_change = new_size - old_size
    
    conn.execute("UPDATE files SET size = ? WHERE id = ?", (new_size, file_id))
    
    # We need the owner and workspace of the file
    file_info = conn.execute("SELECT uploaded_by, workspace_id FROM files WHERE id = ?", (file_id,)).fetchone()
    update_quota(file_info["uploaded_by"], file_info["workspace_id"], size_change, conn)
    
    conn.commit()
    conn.close()
    
    return {"ok": True}

# ============ File Extractors ============

def extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".txt":
        with open(file_path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext in (".csv", ".tsv"):
        try:
            df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
            return f"CSV file, {len(df)} rows x {len(df.columns)} cols\nColumns: {', '.join(df.columns)}\n\n{df.head(100).to_string(index=False)}"
        except Exception:
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                return f.read()
    if ext == ".json":
        with open(file_path, encoding="utf-8") as f:
            data = json.load(f)
        return json.dumps(data, ensure_ascii=False, indent=2)
    if ext in (".xlsx", ".xls"):
        wb = load_workbook(file_path, read_only=True)
        sheets_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(str(row))
            sheets_text.append(f"=== {sheet_name} ({len(rows)}琛? ===\n" + "\n".join(rows[:200]))
        wb.close()
        return "\n\n".join(sheets_text)
    if ext == ".docx":
        doc = Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    with open(file_path, encoding="utf-8", errors="ignore") as f:
        return f.read()

# ============ Analysis ============

def analyze_text(content: str, file_type: str) -> dict:
    lines = content.split("\n")
    non_empty = [l for l in lines if l.strip()]
    words = re.findall(r"[\w\u4e00-\u9fff]+", content)

    stop = {"the", "a", "an", "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "do", "does", "did", "will", "would", "could", "should", "may", "might", "must", "can", "this", "that", "these", "those", "it", "its", "of", "in", "to", "for", "with", "on", "at", "by", "from", "as", "into", "through", "during", "before", "after", "and", "but", "or", "nor", "not", "so", "yet", "both", "either", "neither", "each", "every", "all", "any", "few", "more", "most", "other", "some", "such", "no", "only", "own", "same", "than", "too", "very"}
    filtered = [w for w in words if w.lower() not in stop and len(w) > 1]
    top_words = Counter(filtered).most_common(30)

    numbers = re.findall(r"-?\d+\.?\d*", content)
    if numbers:
        nums = [float(n) for n in numbers if abs(float(n)) < 1e15]
        nums_stats = {
            "count": len(nums),
            "min": round(min(nums), 4),
            "max": round(max(nums), 4),
            "avg": round(sum(nums) / len(nums), 4),
        } if nums else None
    else:
        nums_stats = None

    return {
        "char_count": len(content),
        "line_count": len(lines),
        "non_empty_lines": len(non_empty),
        "word_count": len(words),
        "top_words": top_words,
        "number_stats": nums_stats,
        "preview": content[:500],
    }

def _dedup_columns(raw_header: tuple) -> list:
    """Convert xlsx header tuple to guaranteed-unique string column names."""
    seen: dict = {}
    result = []
    for i, h in enumerate(raw_header, 1):
        name = str(h).strip() if h is not None else ""
        base = name if name else f"列{i}"
        if base in seen:
            seen[base] += 1
            result.append(f"{base}_{seen[base]}")
        else:
            seen[base] = 0
            result.append(base)
    return result

def analyze_csv_table(file_path: Path) -> dict:
    ext = file_path.suffix.lower()
    try:
        if ext == ".csv":
            df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
        else:
            wb = load_workbook(file_path, read_only=True)
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            wb.close()
            if not rows:
                return {"rows": 0, "columns": 0, "column_info": [], "head": []}
            header = _dedup_columns(rows[0])
            n = len(header)
            data_rows = []
            for row in rows[1:]:
                cells = list(row)
                if len(cells) < n:
                    cells += [""] * (n - len(cells))
                data_rows.append([str(c) if c is not None else "" for c in cells[:n]])
            df = pd.DataFrame(data_rows, columns=header)

        if df.empty:
            return {"rows": 0, "columns": len(df.columns), "column_info": [], "head": []}

        cols_info = []
        for i, col in enumerate(df.columns):
            try:
                col_series = df.iloc[:, i]
                series = pd.to_numeric(col_series, errors="coerce")
                info = {
                    "name": str(col),
                    "type": str(col_series.dtype),
                    "null_count": int(col_series.isna().sum()),
                    "unique_count": int(col_series.nunique()),
                }
                if series.notna().any():
                    info["numeric_stats"] = {
                        "min": round(float(series.min()), 4),
                        "max": round(float(series.max()), 4),
                        "mean": round(float(series.mean()), 4),
                    }
            except Exception:
                info = {"name": str(col), "type": "object", "null_count": 0, "unique_count": 0}
            cols_info.append(info)

        try:
            head_dict = df.head(20).to_dict(orient="records")
        except Exception:
            head_dict = []

        return {
            "rows": len(df),
            "columns": len(df.columns),
            "column_info": cols_info,
            "head": head_dict,
        }
    except Exception as e:
        return {"error": f"瑙ｆ瀽澶辫触: {str(e)}"}

# ============ File API (all require auth) ============

@app.get("/api/files")
async def api_list_files(folder: str = "", workspace_id: str = "", current_user: dict = Depends(get_user)):
    conn = get_db()
    all_rows = conn.execute("SELECT * FROM files").fetchall()
    fav_rows = conn.execute("SELECT file_id FROM file_favorites WHERE username = ?", (current_user["username"],)).fetchall()
    fav_ids = {r["file_id"] for r in fav_rows}
    result = []
    for r in all_rows:
        if folder and r["folder"] != folder:
            continue
        ws = r["workspace_id"] if "workspace_id" in r.keys() else ""
        if workspace_id and ws != workspace_id:
            perms = get_file_perms(r["id"], conn)
            if perms.get("visibility") != "public":
                continue
        access = resolve_file_access(r["id"], ws, r["uploaded_by"], current_user, conn)
        if not access["visible"] or not access["can_view"]:
            continue
        d = file_to_dict(r)
        d["favorited"] = d["id"] in fav_ids
        d["permissions"] = get_file_perms(r["id"], conn)
        d["access"] = access
        result.append(d)
    conn.close()
    return result


@app.get("/api/files/page")
async def api_list_files_page(
    folder: str = "",
    workspace_id: str = "",
    q: str = "",
    file_type: str = "",
    sort_by: str = "uploaded_at",
    sort_order: str = "desc",
    limit: int = 50,
    offset: int = 0,
    current_user: dict = Depends(get_user),
):
    """Paginated file list with filtering/sorting, compatible with existing file card rendering."""
    limit = max(1, min(limit, 200))
    offset = max(0, offset)
    sort_by_map = {
        "uploaded_at": "uploaded_at",
        "filename": "filename",
        "size": "size",
        "view_count": "view_count",
    }
    order_col = sort_by_map.get(sort_by, "uploaded_at")
    order_dir = "ASC" if str(sort_order).lower() == "asc" else "DESC"

    conn = get_db()
    fav_rows = conn.execute(
        "SELECT file_id FROM file_favorites WHERE username = ?",
        (current_user["username"],),
    ).fetchall()
    fav_ids = {r["file_id"] for r in fav_rows}

    where_parts = []
    params = []
    if folder:
        where_parts.append("folder = ?")
        params.append(folder)
    # workspace filter is handled after access resolution so public files from
    # other workspaces can still appear in the main list.
    if q:
        where_parts.append("filename LIKE ?")
        params.append(f"%{q.strip()}%")
    if file_type:
        where_parts.append("type = ?")
        params.append(file_type if str(file_type).startswith(".") else f".{file_type}")

    where_sql = f"WHERE {' AND '.join(where_parts)}" if where_parts else ""
    sql = f"SELECT * FROM files {where_sql} ORDER BY {order_col} {order_dir}"
    all_rows = conn.execute(sql, params).fetchall()

    visible_rows = []
    for r in all_rows:
        ws = r["workspace_id"] if "workspace_id" in r.keys() else ""
        if workspace_id and ws != workspace_id:
            perms = get_file_perms(r["id"], conn)
            if perms.get("visibility") != "public":
                continue
        access = resolve_file_access(r["id"], ws, r["uploaded_by"], current_user, conn)
        if not access["visible"] or not access["can_view"]:
            continue
        d = file_to_dict(r)
        d["favorited"] = d["id"] in fav_ids
        d["permissions"] = get_file_perms(r["id"], conn)
        d["access"] = access
        visible_rows.append(d)

    total = len(visible_rows)
    page_items = visible_rows[offset: offset + limit]
    conn.close()
    return {
        "items": page_items,
        "total": total,
        "limit": limit,
        "offset": offset,
        "has_more": offset + limit < total,
    }

@app.get("/api/files/meta")
async def api_list_files_meta(workspace_id: str = "", current_user: dict = Depends(get_user)):
    """Lightweight endpoint returning only file metadata for the sidebar tree (no analysis)."""
    conn = get_db()
    rows = conn.execute("SELECT id, filename, size, type, folder, workspace_id, view_count, uploaded_at, uploaded_by FROM files ORDER BY filename").fetchall()
    fav_rows = conn.execute("SELECT file_id FROM file_favorites WHERE username = ?", (current_user["username"],)).fetchall()
    fav_ids = {r["file_id"] for r in fav_rows}
    result = []
    for r in rows:
        ws = r["workspace_id"] if "workspace_id" in r.keys() else ""
        if workspace_id and ws != workspace_id:
            perms = get_file_perms(r["id"], conn)
            if perms.get("visibility") != "public":
                continue
        access = resolve_file_access(r["id"], ws, r["uploaded_by"], current_user, conn)
        if not access["visible"] or not access["can_view"]:
            continue
        d = dict(r)
        d["favorited"] = d["id"] in fav_ids
        d["access"] = access
        result.append(d)
    conn.close()
    return result

@app.get("/api/file-links")
async def api_list_file_links(current_user: dict = Depends(get_user)):
    conn = get_db()
    file_rows = conn.execute(
        "SELECT id, filename, size, type, folder, workspace_id, view_count, uploaded_at, uploaded_by, analysis FROM files ORDER BY filename"
    ).fetchall()
    visible_files = {}
    nodes = []
    for row in file_rows:
        if not can_user_access_file_row(row, current_user, conn):
            continue
        item = graph_node_to_dict(row)
        visible_files[item["id"]] = item
        nodes.append(item)

    link_rows = conn.execute(
        "SELECT id, file_id, related_file_id, relation_label, created_at, created_by FROM file_links ORDER BY created_at DESC"
    ).fetchall()

    links = []
    for row in link_rows:
        if row["file_id"] not in visible_files or row["related_file_id"] not in visible_files:
            continue
        links.append(dict(row))

    log_rows = conn.execute(
        "SELECT id, link_id, action, file_id, related_file_id, relation_label, created_at, created_by FROM file_link_logs ORDER BY created_at DESC LIMIT 200"
    ).fetchall()

    logs = []
    for row in log_rows:
        if row["file_id"] not in visible_files or row["related_file_id"] not in visible_files:
            continue
        logs.append(dict(row))

    lock_rows = conn.execute(
        "SELECT file_id, username, name, last_heartbeat FROM file_locks"
    ).fetchall()
    now = datetime.now()
    locks = {}
    expired_lock_ids = []
    for row in lock_rows:
        file_id = row["file_id"]
        if file_id not in visible_files:
            continue
        try:
            last_heartbeat = datetime.fromisoformat(row["last_heartbeat"])
        except Exception:
            expired_lock_ids.append(file_id)
            continue
        if now - last_heartbeat > timedelta(seconds=30):
            expired_lock_ids.append(file_id)
            continue
        locks[file_id] = {
            "file_id": file_id,
            "username": row["username"],
            "name": row["name"],
            "last_heartbeat": row["last_heartbeat"],
        }

    if expired_lock_ids:
        conn.executemany("DELETE FROM file_locks WHERE file_id = ?", [(fid,) for fid in expired_lock_ids])
        conn.commit()
    conn.close()

    return {"nodes": nodes, "links": links, "logs": logs, "locks": locks}

@app.post("/api/file-links")
async def api_create_file_link(req: FileLinkCreateRequest, current_user: dict = Depends(get_user)):
    if req.file_id == req.related_file_id:
        raise HTTPException(400, "涓嶈兘鍏宠仈鍚屼竴涓枃浠?")

    file_a, file_b = normalize_file_link_pair(req.file_id, req.related_file_id)
    conn = get_db()
    row_a = conn.execute("SELECT * FROM files WHERE id = ?", (file_a,)).fetchone()
    row_b = conn.execute("SELECT * FROM files WHERE id = ?", (file_b,)).fetchone()
    if not row_a or not row_b:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    if not can_user_access_file_row(row_a, current_user, conn) or not can_user_access_file_row(row_b, current_user, conn):
        conn.close()
        raise HTTPException(403, "鏃犳潈璁块棶鍏宠仈鏂囦欢")

    link_id = uuid.uuid4().hex[:16]
    now = datetime.now().isoformat()
    conn.execute(
        """
        INSERT OR REPLACE INTO file_links
        (id, file_id, related_file_id, relation_label, created_at, created_by)
        VALUES (
            COALESCE((SELECT id FROM file_links WHERE file_id = ? AND related_file_id = ?), ?),
            ?, ?, ?, ?, ?
        )
        """,
        (file_a, file_b, link_id, file_a, file_b, req.relation_label.strip(), now, current_user["username"])
    )
    saved = conn.execute(
        "SELECT id, file_id, related_file_id, relation_label, created_at, created_by FROM file_links WHERE file_id = ? AND related_file_id = ?",
        (file_a, file_b)
    ).fetchone()
    write_file_link_log(saved["id"], "create", file_a, file_b, saved["relation_label"], current_user["username"], conn)
    conn.commit()
    conn.close()
    return dict(saved)

@app.delete("/api/file-links/{link_id}")
async def api_delete_file_link(link_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    row = conn.execute("SELECT * FROM file_links WHERE id = ?", (link_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鍏宠仈涓嶅瓨鍦?")

    row_a = conn.execute("SELECT * FROM files WHERE id = ?", (row["file_id"],)).fetchone()
    row_b = conn.execute("SELECT * FROM files WHERE id = ?", (row["related_file_id"],)).fetchone()
    if not row_a or not row_b:
        conn.execute("DELETE FROM file_links WHERE id = ?", (link_id,))
        conn.commit()
        conn.close()
        return {"ok": True}
    if not can_user_access_file_row(row_a, current_user, conn) or not can_user_access_file_row(row_b, current_user, conn):
        conn.close()
        raise HTTPException(403, "鏃犳潈鍒犻櫎璇ュ叧鑱?")

    write_file_link_log(row["id"], "delete", row["file_id"], row["related_file_id"], row["relation_label"], current_user["username"], conn)
    conn.execute("DELETE FROM file_links WHERE id = ?", (link_id,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/api/search")
async def api_search(q: str = "", current_user: dict = Depends(get_user)):
    if not q.strip():
        return []

    conn = get_db()

    # Build permission filter 鈥?returns a set of allowed file_ids (None = all allowed)
    is_admin = current_user["role"] == "admin"
    can_view_others = is_admin or get_permissions(current_user["username"]).get("can_view_others", False)

    def _allowed(file_id: str, uploaded_by: str) -> bool:
        if is_admin or can_view_others:
            return True
        return uploaded_by == current_user["username"]

    results: list[dict] = []

    fts_expr = build_fts_query(q)
    if fts_expr:
        try:
            fts_rows = conn.execute(
                """
                SELECT
                    m.file_id,
                    snippet(files_fts, 0, '<em>', '</em>', '鈥?, 16) AS fn_snip,
                    snippet(files_fts, 1, '<em>', '</em>', '鈥?, 40) AS ct_snip,
                    rank
                FROM files_fts
                JOIN fts_file_map m ON files_fts.rowid = m.fts_rowid
                WHERE files_fts MATCH ?
                ORDER BY rank
                LIMIT 100
                """,
                (fts_expr,),
            ).fetchall()

            file_ids = [r["file_id"] for r in fts_rows]
            if file_ids:
                placeholders = ",".join("?" * len(file_ids))
                file_rows = {
                    r["id"]: r
                    for r in conn.execute(
                        f"SELECT * FROM files WHERE id IN ({placeholders})", file_ids
                    ).fetchall()
                }
                snippets = {r["file_id"]: r for r in fts_rows}

                for fid in file_ids:
                    fr = file_rows.get(fid)
                    if not fr:
                        continue
                    if not _allowed(fid, fr["uploaded_by"]):
                        continue
                    d = file_to_dict(fr)
                    snip = snippets[fid]
                    # Prefer content snippet; fall back to filename snippet
                    raw_snip = snip["ct_snip"] or snip["fn_snip"] or ""
                    d["snippet"] = raw_snip
                    d["match_type"] = "content" if snip["ct_snip"] else "filename"
                    results.append(d)
        except Exception as e:
            print(f"FTS search error: {e}")
            # Fall through to LIKE fallback

    # LIKE fallback: runs when FTS returned nothing or failed
    if not results:
        like_q = f"%{q.strip()}%"
        if is_admin or can_view_others:
            sql = "SELECT * FROM files WHERE (filename LIKE ? OR analysis LIKE ?)"
            params: tuple = (like_q, like_q)
        else:
            sql = "SELECT * FROM files WHERE uploaded_by = ? AND (filename LIKE ? OR analysis LIKE ?)"
            params = (current_user["username"], like_q, like_q)
        for row in conn.execute(sql, params).fetchall():
            d = file_to_dict(row)
            d["snippet"] = ""
            d["match_type"] = "filename"
            results.append(d)

    conn.close()
    return results


def _get_settings_llm(settings: dict) -> tuple:
    """Load active LLM settings and return (api_key, model, base_url, display_name)."""
    llm_configs_str = settings.get("llm_configs", "[]")
    active_idx = int(settings.get("active_llm_index", "0"))
    try:
        configs = json.loads(llm_configs_str)
        if configs and 0 <= active_idx < len(configs):
            conf = configs[active_idx]
            model = (conf.get("llm_model") or "").strip()
            return (
                (conf.get("llm_api_key") or "").strip(),
                model,
                (conf.get("llm_base_url") or "").strip(),
                (conf.get("name") or model or "?????")
            )
    except Exception:
        pass
    model = settings.get("llm_model", "").strip()
    return (
        settings.get("llm_api_key", "").strip(),
        model,
        settings.get("llm_base_url", "https://api.openai.com/v1").strip(),
        model or "gpt-3.5-turbo"
    )


def _get_file_content_from_fts(file_id: str, conn) -> str:
    """Read full file text from FTS table by file_id."""
    try:
        row = conn.execute(
            "SELECT f.content FROM files_fts f JOIN fts_file_map m ON f.rowid = m.fts_rowid WHERE m.file_id = ?",
            (file_id,)
        ).fetchone()
        return (row[0] or "") if row else ""
    except Exception:
        return ""


def _extract_relevant_chunk(content: str, question: str, chunk_size: int = 900) -> str:
    """Find a relevant chunk in content for the given question."""
    if not content:
        return ""
    if len(content) <= chunk_size:
        return content.strip()
    keywords = list(set(w.lower() for w in re.findall(r"[\w\u4e00-\u9fff]{2,}", question)))
    content_lower = content.lower()
    best_pos, best_score = 0, -1
    step = max(1, chunk_size // 4)
    for start in range(0, len(content) - chunk_size + 1, step):
        window = content_lower[start:start + chunk_size]
        score = sum(window.count(kw) for kw in keywords)
        if score > best_score:
            best_score, best_pos = score, start
    chunk = content[best_pos:best_pos + chunk_size].strip()
    prefix = "..." if best_pos > 0 else ""
    suffix = "..." if best_pos + chunk_size < len(content) else ""
    return prefix + chunk + suffix


def _format_kb_answer(snippets: list) -> str:
    """Format snippet list into a readable markdown answer."""
    parts = []
    for s in snippets:
        content = (s.get("content") or "").strip()
        if content:
            parts.append(f"**{s.get('filename', '?????')}**\n\n{content}")
    if not parts:
        return "????????"
    sep = "\n\n---\n\n"
    return "???????????????\n\n" + sep.join(parts)


@app.get("/api/settings")
async def get_system_settings(authorization: Optional[str] = Header(None)):
    user = get_admin_user(authorization)
    conn = get_db()
    conn.execute("CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT NOT NULL)")
    rows = conn.execute("SELECT key, value FROM settings").fetchall()
    conn.close()
    settings = {r["key"]: r["value"] for r in rows}
    return {
        "llm_api_key": settings.get("llm_api_key", ""),
        "llm_model": settings.get("llm_model", "gpt-3.5-turbo"),
        "llm_base_url": settings.get("llm_base_url", "https://api.openai.com/v1"),
        "llm_configs": settings.get("llm_configs", "[]"),
        "active_llm_index": int(settings.get("active_llm_index", "0")),
    }

@app.put("/api/settings")
async def update_system_settings(req: SettingsUpdate, authorization: Optional[str] = Header(None)):
    user = get_admin_user(authorization)
    conn = get_db()
    conn.execute("CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT NOT NULL)")
    for k, v in req.dict(exclude_none=True).items():
        val = str(v) if v is not None else ""
        if isinstance(v, bool):
            val = "1" if v else "0"
        conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", (k, val))
    conn.commit()
    conn.close()
    return {"ok": True}

def _get_report_db_path(conn) -> str:
    row = conn.execute("SELECT value FROM settings WHERE key = 'report_db_path'").fetchone()
    if row and row["value"] and row["value"].strip():
        return row["value"].strip()
    return "backend/data.db"

def _resolve_report_db_path(path_str: str) -> Path:
    p = Path(path_str)
    if not p.is_absolute():
        p = (BASE_DIR / p).resolve()
    return p

@app.get("/api/reports/access")
async def report_access(current_user: dict = Depends(get_user)):
    if current_user["role"] == "admin":
        return {"visible": True}
    perms = get_permissions(current_user["username"])
    return {"visible": perms.get("can_view_report", False)}

@app.get("/api/reports/tables")
async def report_tables(current_user: dict = Depends(get_user)):
    conn = get_db()
    db_path_str = _get_report_db_path(conn)
    conn.close()
    if current_user["role"] != "admin" and not get_permissions(current_user["username"]).get("can_view_report", False):
        raise HTTPException(403, "鏃犳潈璁块棶鎶ヨ〃")
    db_path = _resolve_report_db_path(db_path_str)
    if not db_path.exists():
        raise HTTPException(404, f"鎶ヨ〃鏁版嵁搴撲笉瀛樺湪: {db_path}")
    rep = sqlite3.connect(str(db_path))
    rep.row_factory = sqlite3.Row
    try:
        rows = rep.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"
        ).fetchall()
        return {"tables": [r["name"] for r in rows], "db_path": str(db_path)}
    finally:
        rep.close()

@app.get("/api/reports/data")
async def report_table_data(table: str, limit: int = 200, current_user: dict = Depends(get_user)):
    conn = get_db()
    db_path_str = _get_report_db_path(conn)
    conn.close()
    if current_user["role"] != "admin" and not get_permissions(current_user["username"]).get("can_view_report", False):
        raise HTTPException(403, "鏃犳潈璁块棶鎶ヨ〃")
    if not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", table):
        raise HTTPException(400, "鏃犳晥琛ㄥ悕")
    limit = max(1, min(limit, 1000))
    db_path = _resolve_report_db_path(db_path_str)
    if not db_path.exists():
        raise HTTPException(404, f"鎶ヨ〃鏁版嵁搴撲笉瀛樺湪: {db_path}")
    rep = sqlite3.connect(str(db_path))
    rep.row_factory = sqlite3.Row
    try:
        exists = rep.execute(
            "SELECT 1 FROM sqlite_master WHERE type='table' AND name = ?",
            (table,),
        ).fetchone()
        if not exists:
            raise HTTPException(404, "鏁版嵁琛ㄤ笉瀛樺湪")
        cnt = rep.execute(f'SELECT COUNT(*) AS cnt FROM "{table}"').fetchone()["cnt"]
        rows = rep.execute(f'SELECT * FROM "{table}" LIMIT ?', (limit,)).fetchall()
        data = [dict(r) for r in rows]
        columns = list(data[0].keys()) if data else []
        return {"table": table, "columns": columns, "rows": data, "total": cnt, "limit": limit}
    finally:
        rep.close()

@app.post("/api/settings/test-llm")
async def test_llm_settings(req: LlmTestRequest, current_user: dict = Depends(get_admin_user)):
    api_key = (req.llm_api_key or "").strip()
    model = (req.llm_model or "").strip()
    base_url = (req.llm_base_url or "").strip()
    if not api_key or not model or not base_url:
        raise HTTPException(400, "llm_api_key / llm_model / llm_base_url 涓嶈兘涓虹┖")

    import httpx
    is_anthropic = _is_anthropic_protocol(base_url, model)
    if is_anthropic:
        url = base_url.rstrip("/") + "/messages"
        payload_data = {
            "model": model,
            "messages": [{"role": "user", "content": "Reply with OK only."}],
            "max_tokens": 16,
        }
        headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "Content-Type": "application/json",
        }
    else:
        url = base_url.rstrip("/") + "/chat/completions"
        payload_data = {
            "model": model,
            "messages": [{"role": "user", "content": "Reply with OK only."}],
            "max_tokens": 8
        }
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(url, json=payload_data, headers=headers, timeout=30.0)
            resp.raise_for_status()
        return {"ok": True, "message": "杩炴帴鎴愬姛"}
    except Exception as e:
        raise HTTPException(400, f"杩炴帴澶辫触: {e}")

@app.get("/api/qa/model-info")
async def get_qa_model_info(current_user: dict = Depends(get_user)):
    conn = get_db()
    conn.execute("CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT NOT NULL)")
    rows = conn.execute("SELECT key, value FROM settings").fetchall()
    conn.close()
    s = {r["key"]: r["value"] for r in rows}
    _, model, _, name = _get_settings_llm(s)
    return {"model": model, "name": name, "has_llm": bool(model)}


def _load_llm_settings() -> tuple:
    conn = get_db()
    conn.execute("CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT NOT NULL)")
    rows = conn.execute("SELECT key, value FROM settings").fetchall()
    conn.close()
    return _get_settings_llm({r["key"]: r["value"] for r in rows})


def _is_anthropic_protocol(base_url: str, model: str) -> bool:
    b = (base_url or "").lower()
    m = (model or "").lower()
    return ("anthropic" in b) or m.startswith("claude")


def _extract_anthropic_text(data: dict) -> str:
    parts = data.get("content") or []
    texts = []
    for p in parts:
        if isinstance(p, dict) and p.get("type") == "text":
            texts.append(p.get("text", ""))
    return "".join(texts).strip()


def _build_qa_snippets(search_results: list, question: str) -> list:
    """Build QA snippets from search results."""
    def _strip_tags(text: str) -> str:
        return re.sub(r"<[^>]+>", "", text or "").replace("?", " ").strip()

    snippets = []
    conn = get_db()
    try:
        for item in (search_results or [])[:5]:
            short_snip = _strip_tags(item.get("snippet") or "")
            file_id = item.get("id")
            larger_content = short_snip
            if file_id:
                full = _get_file_content_from_fts(file_id, conn)
                if full:
                    larger_content = _extract_relevant_chunk(full, question, 900)
            if not larger_content:
                continue
            snippets.append({
                "file_id": file_id,
                "filename": item.get("filename"),
                "type": item.get("type"),
                "snippet": (short_snip or larger_content)[:200],
                "content": larger_content,
            })
    finally:
        conn.close()
    return snippets


@app.get("/api/files/favorites")
async def api_list_favorites(current_user: dict = Depends(get_user)):
    """List all favorited files for the current user."""
    conn = get_db()
    rows = conn.execute(
        """
        SELECT f.* FROM files f
        INNER JOIN file_favorites fav ON f.id = fav.file_id
        WHERE fav.username = ?
        ORDER BY fav.created_at DESC
        """,
        (current_user["username"],),
    ).fetchall()
    result = []
    for r in rows:
        ws = r["workspace_id"] if "workspace_id" in r.keys() else ""
        access = resolve_file_access(r["id"], ws, r["uploaded_by"], current_user, conn)
        if not access["visible"] or not access["can_view"]:
            continue
        d = file_to_dict(r)
        d["favorited"] = True
        d["permissions"] = get_file_perms(r["id"], conn)
        d["access"] = access
        result.append(d)
    conn.close()
    return result


@app.get("/api/files/{file_id}")
async def api_get_file(file_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    ws = row["workspace_id"] if "workspace_id" in row.keys() else ""
    access = resolve_file_access(file_id, ws, row["uploaded_by"], current_user, conn)
    if not access["visible"] or not access["can_view"]:
        conn.close()
        raise HTTPException(403, "鏃犳潈闄愭煡鐪嬫鏂囦欢")
    f = file_to_dict(row)
    f["permissions"] = get_file_perms(file_id, conn)
    f["access"] = access
    c.execute("UPDATE files SET view_count = view_count + 1 WHERE id = ?", (file_id,))
    conn.commit()
    f["view_count"] = f["view_count"] + 1
    conn.close()
    return f

@app.patch("/api/files/{file_id}/move")
async def api_move_file(file_id: str, req: dict, current_user: dict = Depends(get_user)):
    target_folder = req.get("folder")
    target_workspace = req.get("workspace_id")
    
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT uploaded_by, workspace_id FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    
    # Permission check for source file
    is_owner = row["uploaded_by"] == current_user["username"]
    if current_user["role"] != "admin" and not is_owner:
        perms = get_permissions(current_user["username"])
        if not perms.get("can_edit_others", False):
            # Check if user is admin of the current workspace
            ws_id = row["workspace_id"]
            if ws_id:
                ws_member = conn.execute("SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?", (ws_id, current_user["username"])).fetchone()
                if not ws_member or ws_member["role"] != "admin":
                    conn.close()
                    raise HTTPException(403, "鏃犳潈绉诲姩姝ゆ枃浠?")
            else:
                conn.close()
                raise HTTPException(403, "鏃犳潈绉诲姩姝ゆ枃浠?")

    # Permission check for target workspace
    if target_workspace is not None and target_workspace != "":
        if current_user["role"] != "admin":
            ws_member = conn.execute("SELECT role FROM workspace_members WHERE workspace_id = ? AND username = ?", (target_workspace, current_user["username"])).fetchone()
            if not ws_member:
                conn.close()
                raise HTTPException(403, "鎮ㄤ笉鏄洰鏍囧伐浣滅┖闂寸殑鎴愬憳锛屾棤娉曠Щ鍔ㄦ枃浠跺埌璇ョ┖闂?")

    updates = []
    params = []
    if target_folder is not None:
        updates.append("folder = ?")
        params.append(target_folder.strip().strip("/"))
    if target_workspace is not None:
        updates.append("workspace_id = ?")
        params.append(target_workspace)
    
    if updates:
        # Calculate total size of file + versions
        total_size = row["size"] if "size" in row.keys() else 0
        if not total_size:
            # Re-fetch with all columns
            fr = conn.execute("SELECT size FROM files WHERE id = ?", (file_id,)).fetchone()
            total_size = fr["size"]
        
        v_rows = conn.execute("SELECT size FROM file_versions WHERE file_id = ?", (file_id,)).fetchall()
        for vr in v_rows:
            total_size += vr["size"]

        old_ws = row["workspace_id"]
        new_ws = target_workspace if target_workspace is not None else old_ws
        
        if old_ws != new_ws:
            # Check quota in new workspace
            if new_ws:
                try:
                    check_quota(row["uploaded_by"], new_ws, total_size, conn)
                except HTTPException:
                    conn.close()
                    raise
            
            # Update quotas
            if old_ws:
                update_quota(row["uploaded_by"], old_ws, -total_size, conn)
            if new_ws:
                update_quota(row["uploaded_by"], new_ws, total_size, conn)

        params.append(file_id)
        c.execute(f"UPDATE files SET {', '.join(updates)} WHERE id = ?", tuple(params))
        conn.commit()
        
    conn.close()
    return {"ok": True}

@app.delete("/api/files/{file_id}")
async def api_delete_file(file_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    f = file_to_dict(row)
    ws = f.get("workspace_id", "")
    access = resolve_file_access(file_id, ws, f["uploaded_by"], current_user, conn)
    if not access["can_delete"]:
        conn.close()
        raise HTTPException(403, "鏃犳潈闄愬垹闄ゆ鏂囦欢")
    # Calculate total size to remove from quota (main file + all versions)
    total_size_to_remove = f["size"]
    versions = conn.execute("SELECT id, size FROM file_versions WHERE file_id = ?", (file_id,)).fetchall()
    for v in versions:
        total_size_to_remove += v["size"]

    fts_delete_file(conn, file_id)
    c.execute("DELETE FROM files WHERE id = ?", (file_id,))
    c.execute("DELETE FROM file_versions WHERE file_id = ?", (file_id,))

    # Update quota
    update_quota(f["uploaded_by"], ws, -total_size_to_remove, conn)
    
    conn.commit()
    conn.close()

    # Clean up versions
    ext = f["type"].lower()
    for v in versions:
        vp = UPLOAD_VERSIONS_DIR / f"{v['id']}{ext}"
        if vp.exists():
            vp.unlink()

    # Clean up main file
    fp = UPLOAD_DIR / f"{file_id}{ext}"
    if fp.exists():
        fp.unlink()
    else:
        # Fallback for older files or mixed extensions
        for e in [".txt", ".csv", ".xlsx", ".xls", ".json", ".tsv", ".md", ".log", ".docx", ".pptx"]:
            fp = UPLOAD_DIR / f"{file_id}{e}"
            if fp.exists():
                fp.unlink()
                break
    return {"ok": True}

@app.post("/api/files")
async def api_create_file(req: FileCreateRequest, current_user: dict = Depends(get_user)):
    file_id = uuid.uuid4().hex[:12]
    ext = req.file_type.lower()
    if ext not in (".txt", ".md", ".docx", ".xlsx", ".csv"):
        raise HTTPException(400, "涓嶆敮鎸佺殑鏂囦欢绫诲瀷")
    safe_name = f"{file_id}{ext}"
    file_path = UPLOAD_DIR / safe_name

    if ext in (".txt", ".md"):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(req.content)
    elif ext == ".docx":
        doc = Document()
        for para in req.content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(str(file_path))
    elif ext in (".xlsx", ".csv"):
        data = json.loads(req.content)
        wb = Workbook()
        ws = wb.active
        for row in data:
            ws.append(row if isinstance(row, list) else [row])
        wb.save(str(file_path))

    file_size = file_path.stat().st_size
    conn = get_db()
    try:
        check_quota(current_user["username"], req.workspace_id, file_size, conn)
    except HTTPException:
        if file_path.exists():
            file_path.unlink()
        conn.close()
        raise

    text_content = extract_text(file_path)
    if ext in (".csv", ".xlsx", ".xls"):
        analysis = analyze_csv_table(file_path)
    else:
        analysis = analyze_text(text_content, ext)

    analysis_json = json.dumps(analysis, ensure_ascii=False)

    c = conn.cursor()
    folder = req.folder
    workspace_id = req.workspace_id
    if folder:
        c.execute("INSERT OR IGNORE INTO folders (folder) VALUES (?)", (folder,))

    display_name = req.filename if req.filename.endswith(ext) else f"{req.filename}{ext}"
    now = datetime.now().isoformat()
    c.execute(
        "INSERT INTO files (id, filename, size, type, folder, workspace_id, view_count, uploaded_at, uploaded_by, analysis) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (file_id, display_name, file_size, ext, folder, workspace_id, 0, now, current_user["username"], analysis_json)
    )
    
    # Update quota
    update_quota(current_user["username"], workspace_id, file_size, conn)

    if req.visibility != "workspace" or not req.allow_view or req.allow_edit or req.allow_delete:
        _write_file_perms(file_id, req.visibility, req.allow_view, req.allow_edit, req.allow_delete, conn)

    fts_index_file(conn, file_id, display_name, ext)
    conn.commit()
    conn.close()

    return {
        "id": file_id,
        "filename": display_name,
        "size": file_size,
        "type": ext,
        "folder": folder,
        "workspace_id": workspace_id,
        "view_count": 0,
        "uploaded_at": now,
        "uploaded_by": current_user["username"],
        "analysis": analysis,
    }

@app.get("/api/files/{file_id}/content")
async def api_get_file_content(file_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    f = file_to_dict(row)
    is_owner = f.get("uploaded_by") == current_user["username"]
    if current_user["role"] != "admin" and not is_owner:
        perms = get_permissions(current_user["username"])
        if not perms.get("can_view_others", False):
            conn.close()
            raise HTTPException(403, "鏃犳潈闄愭煡鐪嬩粬浜烘枃浠?")
    conn.close()
    ext = f["type"].lower()
    file_path = UPLOAD_DIR / f"{file_id}{ext}"
    if not file_path.exists():
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")

    # Check lock and draft
    is_locked, locker = get_file_lock_info(file_id)
    draft_path = UPLOAD_DRAFTS_DIR / f"{file_id}.draft"
    
    if is_locked and draft_path.exists():
        # Prefer draft for real-time viewing/editing
        read_path = draft_path
    else:
        read_path = file_path

    if ext in (".txt", ".md", ".html", ".json", ".xml", ".js", ".css", ".log", ".yaml", ".yml", ".py", ".java", ".c", ".cpp", ".h", ".sh", ".bat", ".ps1", ".sql", ".go", ".rs", ".rb", ".php", ".ts", ".tsx", ".jsx"):
        with open(read_path, encoding="utf-8", errors="ignore") as fp:
            return {"content": fp.read(), "locked_by": locker if is_locked else None}
    elif ext == ".docx":
        # Only return static text for docx in simple editor
        doc = Document(str(read_path))
        text = "\n\n".join(p.text for p in doc.paragraphs)
        return {"content": text, "locked_by": locker if is_locked else None}
    elif ext in (".xlsx", ".xls"):
        # Table data might not have a .draft equivalent in this simplified sync, 
        # but we follow the same logic.
        wb = load_workbook(read_path, read_only=True)
        ws = wb.active
        data = []
        for row in ws.iter_rows(values_only=True):
            data.append([str(c) if c is not None else "" for c in row])
        wb.close()
        return {"content": data, "locked_by": locker if is_locked else None}
    elif ext == ".csv":
        import csv
        with open(read_path, encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            data = [row for row in reader]
        return {"content": data, "locked_by": locker if is_locked else None}
    raise HTTPException(400, "姝ゆ枃浠剁被鍨嬩笉鏀寔缂栬緫")

@app.get("/api/files/{file_id}/download")
async def api_download_file(file_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    f = file_to_dict(row)
    is_owner = f.get("uploaded_by") == current_user["username"]
    if current_user["role"] != "admin" and not is_owner:
        conn.close()
        raise HTTPException(403, "鏃犳潈闄愪笅杞戒粬浜烘枃浠?")
    conn.close()
    ext = f["type"].lower()
    file_path = UPLOAD_DIR / f"{file_id}{ext}"
    if not file_path.exists():
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    from fastapi.responses import FileResponse
    return FileResponse(
        path=file_path,
        filename=f["filename"],
        media_type="application/octet-stream",
    )

@app.put("/api/files/{file_id}")
async def api_update_file(file_id: str, req: FileUpdateRequest, current_user: dict = Depends(get_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    f = file_to_dict(row)
    old_size = f["size"]
    ws = f.get("workspace_id", "")
    access = resolve_file_access(file_id, ws, f["uploaded_by"], current_user, conn)
    if not access["can_edit"]:
        conn.close()
        raise HTTPException(403, "鏃犳潈闄愮紪杈戞鏂囦欢")
    
    ext = f["type"].lower()
    file_path = UPLOAD_DIR / f"{file_id}{ext}"
    if not file_path.exists():
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")

    # Before overwriting, save a version
    # Note: save_file_version will increment quota for the new version
    save_file_version(file_id, current_user["username"])

    if ext in (".txt", ".md", ".html", ".js", ".css", ".json", ".xml", ".yaml", ".yml", ".py", ".sql", ".sh", ".bat"):
        with open(file_path, "w", encoding="utf-8") as fp:
            fp.write(req.content)
    elif ext == ".docx":
        doc = Document()
        for para in req.content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(str(file_path))
    elif ext in (".xlsx", ".csv"):
        data = json.loads(req.content)
        wb = Workbook()
        ws_wb = wb.active
        for r in data:
            ws_wb.append(r if isinstance(r, list) else [r])
        wb.save(str(file_path))
    else:
        conn.close()
        raise HTTPException(400, "姝ゆ枃浠剁被鍨嬩笉鏀寔缂栬緫")
    
    new_size = file_path.stat().st_size
    size_change = new_size - old_size
    
    # Check if the change exceeds quota
    # If size_change is negative, it's fine. If positive, check.
    if size_change > 0:
        try:
            check_quota(f["uploaded_by"], ws, size_change, conn)
        except HTTPException:
            # Note: We already saved a version, so we could technically roll back,
            # but for now we just raise the error. The file is already updated on disk.
            # Ideally we should check quota BEFORE saving, but we don't know the new size yet for docx/xlsx.
            # For simplicity, we update the quota and allow it this time if it's already on disk,
            # or we could try to revert. 
            pass 

    # Update current file metadata
    conn.execute(
        "UPDATE files SET size = ?, last_modified_at = ?, last_modified_by = ? WHERE id = ?",
        (new_size, datetime.now().isoformat(), current_user["username"], file_id),
    )
    update_quota(f["uploaded_by"], ws, size_change, conn)
    
    # Release lock
    conn.execute("DELETE FROM file_locks WHERE file_id = ? AND username = ?", (file_id, current_user["username"]))

    fts_index_file(conn, file_id, f["filename"], ext)
    conn.commit()
    conn.close()
    
    draft_path = UPLOAD_DRAFTS_DIR / f"{file_id}.draft"
    if draft_path.exists():
        draft_path.unlink()

    return {"ok": True, "size": new_size}

# ============ Folders API ============

class FolderRequest(BaseModel):
    folder: str

@app.get("/api/folders")
async def api_list_folders(current_user: dict = Depends(get_user)):
    conn = get_db()
    rows = conn.execute("SELECT folder FROM folders ORDER BY folder").fetchall()
    conn.close()
    return [r["folder"] for r in rows]

@app.post("/api/folders")
async def api_create_folder(req: FolderRequest, current_user: dict = Depends(get_user)):
    if current_user["role"] != "admin":
        check_perm(current_user, "can_create_folder", is_owner=True)
    folder = req.folder.strip().strip("/")
    if not folder:
        raise HTTPException(400, "璺緞涓嶈兘涓虹┖")
    conn = get_db()
    conn.execute("INSERT OR IGNORE INTO folders (folder) VALUES (?)", (folder,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.delete("/api/folders")
async def api_delete_folder(folder: str, current_user: dict = Depends(get_user)):
    folder = folder.strip().strip("/")
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) as cnt FROM files WHERE folder = ?", (folder,))
    if c.fetchone()["cnt"] > 0:
        conn.close()
        raise HTTPException(400, "璇ヨ矾寰勪笅杩樻湁鏂囦欢锛屾棤娉曞垹闄?")
    conn.execute("DELETE FROM folders WHERE folder = ?", (folder,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.patch("/api/folders/move")
async def api_move_folder(req: dict, current_user: dict = Depends(get_user)):
    """Move/rename a folder: update path prefix for all child folders and files."""
    old_path = req.get("old_path", "").strip().strip("/")
    new_path = req.get("new_path", "").strip().strip("/")
    if not old_path:
        raise HTTPException(400, "鍘熻矾寰勪笉鑳戒负绌?")
    if old_path == new_path:
        raise HTTPException(400, "鐩爣璺緞涓嶈兘涓庡師璺緞鐩稿悓")

    if current_user["role"] != "admin":
        check_perm(current_user, "can_create_folder", is_owner=True)

    conn = get_db()
    c = conn.cursor()

    # Check source folder exists
    c.execute("SELECT folder FROM folders WHERE folder = ?", (old_path,))
    if not c.fetchone():
        conn.close()
        raise HTTPException(404, "鍘熸枃浠跺す涓嶅瓨鍦?")

    # Check target folder doesn't already exist
    if new_path:
        c.execute("SELECT folder FROM folders WHERE folder = ?", (new_path,))
        if c.fetchone():
            conn.close()
            raise HTTPException(400, "鐩爣鏂囦欢澶瑰凡瀛樺湪")
        # Check new_path is not a subfolder of old_path (would create cycle)
        if new_path.startswith(old_path + "/"):
            conn.close()
            raise HTTPException(400, "涓嶈兘灏嗘枃浠跺す绉诲姩鍒板叾鑷韩瀛愮洰褰曚笅")

    try:
        # Rename the folder itself
        if new_path:
            c.execute("UPDATE folders SET folder = ? WHERE folder = ?", (new_path, old_path))
        else:
            c.execute("DELETE FROM folders WHERE folder = ?", (old_path,))

        # Rename all child folders (paths starting with old_path/)
        c.execute("SELECT folder FROM folders WHERE folder LIKE ?", (old_path + "/%",))
        child_folders = [r["folder"] for r in c.fetchall()]
        for cf in child_folders:
            if new_path:
                new_cf = new_path + cf[len(old_path):]
                c.execute("UPDATE folders SET folder = ? WHERE folder = ?", (new_cf, cf))
            else:
                # Moving to root: strip the old_path prefix
                new_cf = cf[len(old_path) + 1:]
                c.execute("UPDATE folders SET folder = ? WHERE folder = ?", (new_cf, cf))

        # Update all files in this folder
        c.execute("UPDATE files SET folder = ? WHERE folder = ?", (new_path, old_path))

        # Update all files in child folders
        for cf in child_folders:
            if new_path:
                new_cf = new_path + cf[len(old_path):]
                c.execute("UPDATE files SET folder = ? WHERE folder = ?", (new_cf, cf))
            else:
                new_cf = cf[len(old_path) + 1:]
                c.execute("UPDATE files SET folder = ? WHERE folder = ?", (new_cf, cf))

        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        raise HTTPException(500, f"绉诲姩澶辫触: {str(e)}")
    conn.close()
    return {"ok": True}

# ============ Favorites API ============

@app.post("/api/files/{file_id}/favorite")
async def api_add_favorite(file_id: str, current_user: dict = Depends(get_user)):
    """Add a file to favorites."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT id FROM files WHERE id = ?", (file_id,))
    if not c.fetchone():
        conn.close()
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    now = datetime.now().isoformat()
    c.execute("INSERT OR IGNORE INTO file_favorites (username, file_id, created_at) VALUES (?, ?, ?)",
              (current_user["username"], file_id, now))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.delete("/api/files/{file_id}/favorite")
async def api_remove_favorite(file_id: str, current_user: dict = Depends(get_user)):
    """Remove a file from favorites."""
    conn = get_db()
    conn.execute("DELETE FROM file_favorites WHERE username = ? AND file_id = ?",
                 (current_user["username"], file_id))
    conn.commit()
    conn.close()
    return {"ok": True}

# ============ Lock & Heartbeat API ============

@app.post("/api/files/{file_id}/lock")
async def api_lock_file(file_id: str, current_user: dict = Depends(get_user)):
    is_locked, locker = get_file_lock_info(file_id)
    if is_locked and locker["username"] != current_user["username"]:
        return {"ok": False, "locker": locker}
    
    now = datetime.now().isoformat()
    conn = get_db()
    conn.execute("INSERT OR REPLACE INTO file_locks (file_id, username, name, last_heartbeat) VALUES (?, ?, ?, ?)",
                 (file_id, current_user["username"], current_user.get("name", current_user["username"]), now))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.post("/api/files/{file_id}/heartbeat")
async def api_file_heartbeat(file_id: str, req: dict, current_user: dict = Depends(get_user)):
    # Verify lock ownership
    is_locked, locker = get_file_lock_info(file_id)
    if not is_locked or locker["username"] != current_user["username"]:
        raise HTTPException(403, "鎮ㄦ湭鎸佹湁璇ユ枃浠剁殑缂栬緫閿?")
    
    # Update heartbeat
    now = datetime.now().isoformat()
    conn = get_db()
    conn.execute("UPDATE file_locks SET last_heartbeat = ? WHERE file_id = ?", (now, file_id))
    conn.commit()
    conn.close()
    
    # Save draft content if provided
    if "content" in req:
        draft_path = UPLOAD_DRAFTS_DIR / f"{file_id}.draft"
        with open(draft_path, "w", encoding="utf-8") as f:
            f.write(req["content"])
            
    return {"ok": True}

@app.delete("/api/files/{file_id}/lock")
async def api_unlock_file(file_id: str, current_user: dict = Depends(get_user)):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT username FROM file_locks WHERE file_id = ?", (file_id,))
    row = c.fetchone()
    if row and row["username"] == current_user["username"]:
        conn.execute("DELETE FROM file_locks WHERE file_id = ?", (file_id,))
        # Also clean up draft file
        draft_path = UPLOAD_DRAFTS_DIR / f"{file_id}.draft"
        if draft_path.exists():
            draft_path.unlink()
    conn.commit()
    conn.close()
    return {"ok": True}

# ============ Dashboard API ============

@app.get("/api/dashboard")
async def api_dashboard(
    folder: str = "",
    workspace_id: str = "",
    favorites_only: bool = False,
    current_user: dict = Depends(get_user),
):
    conn = get_db()
    fav_ids = set()
    if favorites_only:
        fav_rows = conn.execute(
            "SELECT file_id FROM file_favorites WHERE username = ?",
            (current_user["username"],),
        ).fetchall()
        fav_ids = {r["file_id"] for r in fav_rows}
    rows = conn.execute("SELECT * FROM files").fetchall()
    files = []
    for r in rows:
        if favorites_only and r["id"] not in fav_ids:
            continue
        if folder and r["folder"] != folder:
            continue
        ws = r["workspace_id"] if "workspace_id" in r.keys() else ""
        if workspace_id and ws != workspace_id:
            perms = get_file_perms(r["id"], conn)
            if perms.get("visibility") != "public":
                continue
        access = resolve_file_access(r["id"], ws, r["uploaded_by"], current_user, conn)
        if not access["visible"] or not access["can_view"]:
            continue
        files.append(file_to_dict(r))

    total_files = len(files)
    total_size = sum(f.get("size", 0) for f in files)
    file_types = {}
    for f in files:
        t = f.get("type", "unknown")
        file_types[t] = file_types.get(t, 0) + 1

    folder_stats = {}
    for f in files:
        folder = f.get("folder", "") or "/"
        folder_stats[folder] = folder_stats.get(folder, 0) + 1

    sorted_by_time = sorted(files, key=lambda x: x.get("uploaded_at", ""), reverse=True)
    recent = [{"id": f["id"], "filename": f["filename"], "type": f.get("type", ""), "folder": f.get("folder", ""), "uploaded_at": f.get("uploaded_at", ""), "size": f.get("size", 0)} for f in sorted_by_time[:5]]

    sorted_by_views = sorted(files, key=lambda x: x.get("view_count", 0), reverse=True)
    hot = [{"id": f["id"], "filename": f["filename"], "type": f.get("type", ""), "folder": f.get("folder", ""), "view_count": f.get("view_count", 0)} for f in sorted_by_views[:10] if f.get("view_count", 0) > 0]

    folder_count = len({(f.get("folder", "") or "/") for f in files})
    conn.close()

    all_filenames = [f["filename"] for f in files]

    # Aggregate top_words only from Office documents (Word/PowerPoint);
    # code files, plain text, etc. are excluded 鈥?their filenames are used instead.
    OFFICE_TYPES = {".docx", ".pptx"}
    content_word_freq: dict = {}
    for f in files:
        if f.get("type", "") not in OFFICE_TYPES:
            continue
        analysis = f.get("analysis", {})
        if not isinstance(analysis, dict):
            continue
        for entry in (analysis.get("top_words") or []):
            if isinstance(entry, (list, tuple)) and len(entry) == 2:
                word, count = str(entry[0]), int(entry[1])
                content_word_freq[word] = content_word_freq.get(word, 0) + count
    content_words = sorted(content_word_freq.items(), key=lambda x: -x[1])[:120]

    return {
        "total_files": total_files,
        "total_size": total_size,
        "total_folders": folder_count,
        "file_types": file_types,
        "folder_stats": folder_stats,
        "recent_files": recent,
        "hot_files": hot,
        "all_filenames": all_filenames,
        "content_words": content_words,
    }

@app.get("/api/me/stats")
async def api_my_stats(current_user: dict = Depends(get_user)):
    conn = get_db()
    username = current_user["username"]
    
    # User's files
    rows = conn.execute("SELECT * FROM files WHERE uploaded_by = ?", (username,)).fetchall()
    files = [file_to_dict(r) for r in rows]
    
    total_files = len(files)
    total_size = sum(f.get("size", 0) for f in files)
    
    # Top 3 most viewed files by this user
    hot_rows = conn.execute("SELECT filename, view_count, type FROM files WHERE uploaded_by = ? ORDER BY view_count DESC LIMIT 3", (username,)).fetchall()
    hot_files = [dict(r) for r in hot_rows]
    
    # Recent 3 activities 鈥?sort by last modification time if available, else upload time
    recent_rows = conn.execute(
        "SELECT filename, uploaded_at, type, last_modified_at, last_modified_by FROM files WHERE uploaded_by = ? ORDER BY COALESCE(last_modified_at, uploaded_at) DESC LIMIT 3",
        (username,),
    ).fetchall()
    recent_activity = [dict(r) for r in recent_rows]
    
    conn.close()
    
    return {
        "total_files": total_files,
        "total_size": total_size,
        "hot_files": hot_files,
        "recent_activity": recent_activity
    }

# ============ ONLYOFFICE Integration ============

ONLYOFFICE_SERVER = os.environ.get("ONLYOFFICE_SERVER", "http://localhost:8888")
# kb-backend resolves inside the OnlyOffice container to the host via IPv4 only (no IPv6 ambiguity)
ONLYOFFICE_CALLBACK_BASE = os.environ.get("ONLYOFFICE_CALLBACK_BASE", "http://kb-backend:8000")

MIME_TYPES = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".doc": "application/msword",
}

@app.get("/api/files/{file_id}/onlyoffice-download")
async def onlyoffice_download(file_id: str):
    """ONLYOFFICE-only download endpoint (no auth required, used by ONLYOFFICE server)."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    f = file_to_dict(row)
    ext = f["type"].lower()
    file_path = UPLOAD_DIR / f"{file_id}{ext}"
    if not file_path.exists():
        raise HTTPException(404, "鏂囦欢涓嶅瓨鍦?")
    from fastapi.responses import FileResponse
    return FileResponse(
        path=file_path,
        filename=f["filename"],
        media_type=MIME_TYPES.get(ext, "application/octet-stream"),
    )

@app.post("/api/files/{file_id}/onlyoffice-callback")
async def onlyoffice_callback(file_id: str, request: Request):
    """ONLYOFFICE callback - saves edited file back to server."""
    body = await request.json()
    status = body.get("status")
    # status 2 = file saved successfully
    if status == 2:
        download_url = body.get("url")
        if not download_url:
            return {"error": 0, "message": "No URL provided"}
        conn = get_db()
        c = conn.cursor()
        c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
        row = c.fetchone()
        if not row:
            conn.close()
            return {"error": 1, "message": "File not found"}
        f = file_to_dict(row)
        ext = f["type"].lower()
        file_path = UPLOAD_DIR / f"{file_id}{ext}"
        try:
            import httpx
            resp = httpx.get(download_url, follow_redirects=True, timeout=30)
            if resp.status_code == 200:
                # Get locker or users who edited
                users = body.get("users", [])
                saved_by = users[0] if users else "onlyoffice"
                
                # Before saving new content, save a version (ONLY if original exists)
                if file_path.exists():
                    save_file_version(file_id, saved_by)

                with open(file_path, "wb") as fp:
                    fp.write(resp.content)
                new_size = file_path.stat().st_size
                conn.execute(
                    "UPDATE files SET size = ?, last_modified_at = ?, last_modified_by = ? WHERE id = ?",
                    (new_size, datetime.now().isoformat(), saved_by, file_id),
                )
                fts_index_file(conn, file_id, f["filename"], ext)
                conn.commit()
            else:
                print(f"ONLYOFFICE download failed: {resp.status_code}")
        except Exception as e:
            print(f"ONLYOFFICE callback error: {e}")
        finally:
            conn.close()
    # status 3 = force save error
    elif status == 3:
        print(f"ONLYOFFICE force save error for {file_id}")
    return {"error": 0}

@app.get("/api/files/{file_id}/onlyoffice-config")
async def onlyoffice_config(file_id: str, current_user: dict = Depends(get_user)):
    """Return ONLYOFFICE config for opening a file."""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM files WHERE id = ?", (file_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "File not found")
    f = file_to_dict(row)
    ws = row["workspace_id"] if "workspace_id" in row.keys() else ""
    access = resolve_file_access(file_id, ws, row["uploaded_by"], current_user, conn)
    if not access["visible"] or not access["can_view"]:
        conn.close()
        raise HTTPException(403, "No permission to view this file")
    can_edit = bool(access["can_edit"])
    conn.close()

    ext = f["type"].lower()
    if ext == ".docx":
        doc_type = "word"
    elif ext in (".xlsx", ".xls"):
        doc_type = "cell"
    elif ext == ".pptx":
        doc_type = "slide"
    else:
        raise HTTPException(400, "Unsupported file type for ONLYOFFICE")

    download_url = f"{ONLYOFFICE_CALLBACK_BASE}/api/files/{file_id}/onlyoffice-download"
    callback_url = f"{ONLYOFFICE_CALLBACK_BASE}/api/files/{file_id}/onlyoffice-callback"
    doc_key = f"{file_id}_{f['size']}"

    return {
        "onlyoffice_server": ONLYOFFICE_SERVER,
        "document_url": download_url,
        "callback_url": callback_url,
        "document_type": doc_type,
        "file_type": ext.lstrip("."),
        "filename": f["filename"],
        "username": current_user.get("name", current_user["username"]),
        "key": doc_key,
        "can_edit": can_edit,
        "mode": "edit" if can_edit else "view",
    }


# ============ Utility API ============

@app.get("/api/health")
async def api_health():
    return {"ok": True, "service": "file-knowledge-base"}


@app.get("/api/routes")
async def api_routes():
    """Quick route catalog for frontend/debug tooling."""
    data = []
    for r in app.router.routes:
        path = getattr(r, "path", "")
        methods = sorted(list(getattr(r, "methods", []) or []))
        if not path.startswith("/api/"):
            continue
        data.append({"path": path, "methods": methods})
    data.sort(key=lambda x: x["path"])
    return data

# 鈹€鈹€ Serve frontend static pages 鈹€鈹€
FRONTEND_DIR = BASE_DIR / "frontend"

@app.get("/")
async def serve_index():
    return FileResponse(FRONTEND_DIR / "index.html", media_type="text/html")

@app.get("/editor.html")
async def serve_editor():
    return FileResponse(FRONTEND_DIR / "editor.html", media_type="text/html")

@app.get("/graph.html")
async def serve_graph():
    return FileResponse(FRONTEND_DIR / "graph.html", media_type="text/html")

@app.get("/qa.html")
async def serve_qa():
    return FileResponse(FRONTEND_DIR / "qa.html", media_type="text/html")

@app.get("/vue.global.js")
async def serve_vue():
    return FileResponse(FRONTEND_DIR / "vue.global.js", media_type="application/javascript")

@app.get("/marked.min.js")
async def serve_marked():
    return FileResponse(FRONTEND_DIR / "marked.min.js", media_type="application/javascript")

@app.get("/qa.js")
async def serve_qa_js():
    return FileResponse(FRONTEND_DIR / "qa.js", media_type="application/javascript")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)




